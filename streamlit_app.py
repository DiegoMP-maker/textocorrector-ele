"""
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 1: Importaciones y Configuración Base
==================================================================================

Este artefacto contiene:
1. Todas las importaciones de bibliotecas necesarias
2. Configuración base de la aplicación Streamlit
3. Configuración de logging
4. Definición de la versión de la aplicación
"""

import traceback
import streamlit as st
import json
import gspread
import requests
import re
import pandas as pd
import matplotlib.pyplot as plt
import altair as alt
import time
import io
import base64
import numpy as np
from google.oauth2.service_account import Credentials
from datetime import datetime
from openai import OpenAI
from io import BytesIO, StringIO
from PIL import Image
import qrcode
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import logging
from urllib.parse import urlparse
import uuid

# Configuración de logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# Configuración de la página
st.set_page_config(
    page_title="Textocorrector ELE",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Configuración de cache y timeout
st.cache_data.clear()

# Versión de la aplicación
APP_VERSION = "2.1.0"  # Actualizado para la versión reescrita
"""
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 2: Inicialización de Variables y Estados
==================================================================================

Este artefacto contiene:
1. Funciones para inicializar variables de session_state
2. Funciones seguras para acceder y modificar valores del session_state
3. Inicialización del sidebar
"""

# Inicializar variables de session_state si no existen


def init_session_state():
    """
    Inicializa variables de session_state con valores predeterminados seguros
    para evitar KeyError durante la ejecución.
    """
    default_values = {
        "nivel_estudiante": "intermedio",
        "consigna_actual": "",
        "usar_consigna_como_texto": False,
        "texto_correccion_corregir": "",
        "info_adicional_corregir": "",
        "ultima_imagen_url": "",
        "ultima_descripcion": "",
        "ultimo_texto_transcrito": "",
        "tarea_modelo_generada": None,
        "respuesta_modelo_examen": "",
        "inicio_simulacro": None,
        "duracion_simulacro": None,
        "tarea_simulacro": None,
        "simulacro_respuesta_texto": "",
        "request_id": "",
        "usuario_actual": "",
        "correction_result": None,
        "last_correction_time": None,
        "examen_result": None,
        "api_error_count": 0,
        "api_last_error_time": None,
        "circuit_breaker_open": False,
        "ultimo_texto": "",
        "nombre_seleccionado": None,
        "imagen_generada_state": False,
        "imagen_url_state": None,
        "descripcion_state": None,
        "tema_imagen_state": None,
        "descripcion_estudiante_state": "",
        "mostrar_correccion_imagen": False,
        "mostrar_correccion_transcripcion": False,
        "active_tab_index": 0,
        "active_tools_tab_index": 0,
        "tab_navigate_to": None,
        "app_initialized": False
    }

    for key, default_value in default_values.items():
        if key not in st.session_state:
            st.session_state[key] = default_value


# Inicializar session_state
init_session_state()

# Funciones seguras para acceso a session_state


def get_session_var(key, default=None):
    """Obtiene una variable de session_state de forma segura"""
    return st.session_state.get(key, default)


def set_session_var(key, value):
    """Establece una variable en session_state"""
    st.session_state[key] = value


# Generar ID único para esta sesión si no existe
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

# Mensaje de bienvenida en sidebar
st.sidebar.title("📝 Textocorrector ELE")
st.sidebar.info(
    """
    Versión: {0}

    Una herramienta para corrección de textos
    en español con análisis contextual avanzado.

    ID de sesión: {1}
    """.format(APP_VERSION, st.session_state.session_id[:8])
)
"""
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 3: Funciones de Seguridad y Conexión a APIs
==================================================================================

Este artefacto contiene:
1. Funciones para manejo seguro de claves API
2. Implementación del patrón Circuit Breaker para APIs
3. Conexión segura a Google Sheets
4. Configuración segura de clientes API (OpenAI)
5. Funciones de diagnóstico para el estado de conexiones
"""

# --- 1. CONFIGURACIÓN DE CLAVES SEGURAS ---


def get_api_keys():
    """
    Obtiene las claves de API de los secretos de Streamlit con manejo de errores.
    Permite la operación en modo degradado si faltan claves.
    """
    keys = {
        "openai": None,
        "elevenlabs": {"api_key": None, "voice_id": None},
        "google_credentials": None
    }

    try:
        keys["openai"] = st.secrets["OPENAI_API_KEY"]
    except Exception as e:
        logger.warning(f"Error al obtener API Key de OpenAI: {e}")
        st.sidebar.warning(
            "⚠️ API de OpenAI no configurada. Algunas funciones estarán limitadas.")

    try:
        keys["elevenlabs"]["api_key"] = st.secrets["ELEVENLABS_API_KEY"]
        keys["elevenlabs"]["voice_id"] = st.secrets["ELEVENLABS_VOICE_ID"]
    except Exception as e:
        logger.warning(f"Error al obtener configuración de ElevenLabs: {e}")
        st.sidebar.warning(
            "⚠️ API de ElevenLabs no configurada. La función de audio estará deshabilitada.")

    try:
        keys["google_credentials"] = json.loads(
            st.secrets["GOOGLE_CREDENTIALS"])
    except Exception as e:
        logger.warning(f"Error al obtener credenciales de Google: {e}")
        st.sidebar.warning(
            "⚠️ Credenciales de Google no configuradas. El guardado de datos estará deshabilitado.")

    return keys


# Obtener claves de API
api_keys = get_api_keys()

# --- 2. CIRCUIT BREAKER PARA APIS ---


class CircuitBreaker:
    """
    Implementa el patrón Circuit Breaker para APIs externas.
    Previene llamadas repetidas a APIs con fallo.
    """

    def __init__(self, failure_threshold=5, reset_timeout=300):
        self.failure_threshold = failure_threshold  # Número de fallos antes de abrir
        self.reset_timeout = reset_timeout  # Tiempo en segundos antes de reintentar

        # Inicializar contadores para diferentes servicios
        self.services = {
            "openai": {"failures": 0, "last_failure_time": None, "open": False},
            "elevenlabs": {"failures": 0, "last_failure_time": None, "open": False},
            "google_sheets": {"failures": 0, "last_failure_time": None, "open": False}
        }

    def record_failure(self, service_name):
        """Registra un fallo para el servicio especificado"""
        if service_name not in self.services:
            logger.warning(f"Servicio desconocido: {service_name}")
            return

        service = self.services[service_name]
        service["failures"] += 1
        service["last_failure_time"] = time.time()

        if service["failures"] >= self.failure_threshold:
            service["open"] = True
            logger.warning(f"Circuit breaker ABIERTO para {service_name}")

    def record_success(self, service_name):
        """Registra un éxito y restablece contadores para el servicio"""
        if service_name not in self.services:
            return

        service = self.services[service_name]
        service["failures"] = 0
        service["open"] = False

    def can_execute(self, service_name):
        """Determina si se puede ejecutar una llamada al servicio"""
        if service_name not in self.services:
            return True

        service = self.services[service_name]

        # Si el circuit breaker está abierto, verificar si ha pasado el tiempo de reset
        if service["open"]:
            if service["last_failure_time"] is None:
                return True

            elapsed = time.time() - service["last_failure_time"]
            if elapsed > self.reset_timeout:
                # Permitir un reintento
                service["open"] = False
                return True
            else:
                return False

        return True

    def get_status(self):
        """Devuelve el estado actual de todos los servicios"""
        return {name: {"open": info["open"], "failures": info["failures"]}
                for name, info in self.services.items()}


# Inicializar circuit breaker
circuit_breaker = CircuitBreaker()

# --- 3. CONEXIÓN A GOOGLE SHEETS ---


def connect_to_googlesheets():
    """
    Establece conexión con Google Sheets para almacenamiento de datos.
    Retorna un diccionario con los objetos de conexión o None si hay error.
    """
    if api_keys["google_credentials"] is None:
        return None

    if not circuit_breaker.can_execute("google_sheets"):
        st.warning(
            "⚠️ Conexión a Google Sheets temporalmente deshabilitada debido a errores previos.")
        return None

    try:
        scope = [
            "https://spreadsheets.google.com/feeds",
            "https://www.googleapis.com/auth/drive"
        ]

        creds = Credentials.from_service_account_info(
            api_keys["google_credentials"], scopes=scope)
        client_gsheets = gspread.authorize(creds)

        # IDs de los documentos
        CORRECTIONS_DOC_ID = "1GTaS0Bv_VN-wzTq1oiEbDX9_UdlTQXWhC9CLeNHVk_8"
        TRACKING_DOC_ID = "1-OQsMGgWseZ__FyUVh0UtYVOLui_yoTMG0BxxTGPOU8"

        sheets = {}

        # Intentar abrir documentos con manejo de errores para cada uno
        try:
            sheets["corrections"] = client_gsheets.open_by_key(
                CORRECTIONS_DOC_ID).sheet1
            logger.info("Conectado a Historial_Correcciones_ELE")
        except Exception as e:
            logger.error(
                f"Error al conectar con Historial_Correcciones_ELE: {e}")
            sheets["corrections"] = None

        try:
            tracking_doc = client_gsheets.open_by_key(TRACKING_DOC_ID)

            # Verificar si existe la hoja Seguimiento
            try:
                sheets["tracking"] = tracking_doc.worksheet("Seguimiento")
                logger.info("Conectado a hoja Seguimiento")
            except gspread.exceptions.WorksheetNotFound:
                # Crear la hoja si no existe
                sheets["tracking"] = tracking_doc.add_worksheet(
                    title="Seguimiento", rows=100, cols=14)
                # Añadir encabezados a la hoja
                headers = ["Nombre", "Nivel", "Fecha", "Errores Gramática", "Errores Léxico",
                           "Errores Puntuación", "Errores Estructura", "Total Errores",
                           "Puntuación Coherencia", "Puntuación Cohesión", "Puntuación Registro",
                           "Puntuación Adecuación Cultural", "Consejo Final"]
                sheets["tracking"].append_row(headers)
                logger.info("Hoja 'Seguimiento' creada y preparada")
        except Exception as e:
            logger.error(f"Error al conectar con hoja de Seguimiento: {e}")
            sheets["tracking"] = None

        # Verificar si hubo éxito en alguna conexión
        if sheets["corrections"] is not None or sheets["tracking"] is not None:
            circuit_breaker.record_success("google_sheets")
            return sheets
        else:
            circuit_breaker.record_failure("google_sheets")
            return None

    except Exception as e:
        logger.error(f"Error al conectar con Google Sheets: {e}")
        circuit_breaker.record_failure("google_sheets")
        return None


# Establecer conexión con Google Sheets (podría ser None si falla)
sheets_connection = connect_to_googlesheets()

# --- 4. CLIENTE DE OPENAI SEGURO ---


def get_openai_client():
    """
    Crea un cliente de OpenAI con manejo de errores.
    Retorna el cliente o None si no es posible crear la conexión.
    """
    if api_keys["openai"] is None:
        return None

    if not circuit_breaker.can_execute("openai"):
        st.warning(
            "⚠️ Conexión a OpenAI temporalmente deshabilitada debido a errores previos.")
        return None

    try:
        client = OpenAI(api_key=api_keys["openai"])
        circuit_breaker.record_success("openai")
        return client
    except Exception as e:
        logger.error(f"Error al crear cliente OpenAI: {e}")
        circuit_breaker.record_failure("openai")
        return None

# --- 5. UTILIDADES DE DIAGNÓSTICO ---


def show_connection_status():
    """Muestra el estado de conexión de los servicios externos"""
    with st.sidebar.expander("Estado de conexiones", expanded=False):
        status = circuit_breaker.get_status()

        # OpenAI
        if api_keys["openai"] is not None:
            if not status["openai"]["open"]:
                st.sidebar.success("✅ OpenAI: Conectado")
            else:
                st.sidebar.error(
                    f"❌ OpenAI: Desconectado ({status['openai']['failures']} fallos)")
        else:
            st.sidebar.warning("⚠️ OpenAI: No configurado")

        # Google Sheets
        if sheets_connection is not None:
            sheets_status = []
            if sheets_connection["corrections"] is not None:
                sheets_status.append("Historial")
            if sheets_connection["tracking"] is not None:
                sheets_status.append("Seguimiento")

            if sheets_status:
                st.sidebar.success(
                    f"✅ Google Sheets: {', '.join(sheets_status)}")
            else:
                st.sidebar.error("❌ Google Sheets: Error de conexión")
        else:
            if api_keys["google_credentials"] is not None:
                st.sidebar.error("❌ Google Sheets: Error de conexión")
            else:
                st.sidebar.warning("⚠️ Google Sheets: No configurado")

        # ElevenLabs
        if api_keys["elevenlabs"]["api_key"] is not None:
            if not status["elevenlabs"]["open"]:
                st.sidebar.success("✅ ElevenLabs: Conectado")
            else:
                st.sidebar.error(
                    f"❌ ElevenLabs: Desconectado ({status['elevenlabs']['failures']} fallos)")
        else:
            st.sidebar.warning("⚠️ ElevenLabs: No configurado")


# Mostrar estado de conexión en sidebar
show_connection_status()

# --- FUNCIÓN AUXILIAR PARA MANEJO DE EXCEPCIONES ---


def handle_exception(func_name, exception, show_user=True):
    """
    Función de utilidad para manejar excepciones de manera consistente.

    Args:
        func_name: Nombre de la función donde ocurrió el error
        exception: La excepción capturada
        show_user: Si se debe mostrar un mensaje al usuario

    Returns:
        None
    """
    error_msg = f"Error en {func_name}: {str(exception)}"
    logger.error(error_msg)
    logger.error(traceback.format_exc())

    if show_user:
        st.error(f"⚠️ {error_msg}")
        with st.expander("Detalles técnicos", expanded=False):
            st.code(traceback.format_exc())

    return None


"""
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 4: Funciones Core de Procesamiento
==================================================================================

Este artefacto contiene:
1. Funciones para procesamiento de respuestas API (OpenAI, ElevenLabs)
2. Funciones para procesamiento de JSON
3. Integración core con APIs externas 
4. Funciones de procesamiento para corrección y análisis de texto
"""

# --- 1. FUNCIONES DE API DE OPENAI ---


def extract_json_safely(content):
    """
    Extrae contenido JSON de una respuesta con múltiples estrategias.
    Implementa parsing robusto para evitar errores.

    Args:
        content: Contenido de texto que debería contener JSON

    Returns:
        dict: El contenido parseado como JSON o un diccionario con error
    """
    # Si es None o vacío, retornar error inmediatamente
    if not content:
        return {"error": "Contenido vacío, no se puede extraer JSON"}

    # Intento directo
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        # Limpiar el contenido - eliminar caracteres que puedan causar problemas
        # Esto puede ayudar con el problema "unknown extension ?1 at position 13"
        content_clean = re.sub(r'[^\x20-\x7E]', '', content)

        try:
            return json.loads(content_clean)
        except json.JSONDecodeError:
            pass

        # Búsqueda con regex para JSON completo
        # Regex mejorada para JSON anidado
        json_pattern = r'(\{(?:[^{}]|(?1))*\})'
        match = re.search(json_pattern, content_clean, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass

        # Segunda estrategia: buscar cualquier JSON entre llaves
        simple_pattern = r'\{.*\}'
        match = re.search(simple_pattern, content_clean, re.DOTALL)
        if match:
            try:
                # Intentar limpiar el JSON encontrado
                potential_json = match.group(0)
                # Eliminar comillas mal formadas, escape chars, etc.
                potential_json = re.sub(r'[\r\n\t]', ' ', potential_json)
                return json.loads(potential_json)
            except json.JSONDecodeError:
                pass

    # Si no se pudo extraer, devolver un objeto error
    logger.warning(f"No se pudo extraer JSON de: {content[:100]}...")
    return {"error": "No se pudo extraer JSON válido", "raw_content": content[:500]}


def retry_with_backoff(func, max_retries=3, initial_delay=1):
    """
    Ejecuta una función con reintentos y backoff exponencial.

    Args:
        func: Función a ejecutar
        max_retries: Número máximo de reintentos
        initial_delay: Retraso inicial en segundos

    Returns:
        El resultado de la función o levanta la excepción
    """
    for attempt in range(max_retries):
        try:
            return func()
        except (requests.ConnectionError, requests.Timeout) as e:
            # Errores de red específicos - reintentamos
            if attempt == max_retries - 1:
                raise
            delay = initial_delay * (2 ** attempt)  # Backoff exponencial
            logger.info(
                f"Reintento {attempt+1} en {delay} segundos debido a: {str(e)}")
            time.sleep(delay)
        except Exception as e:
            # Otros errores - no reintentamos
            logger.error(f"Error no recuperable: {str(e)}")
            raise


def obtener_json_de_ia(system_msg, user_msg, model="gpt-4-turbo", max_retries=3):
    """
    Obtiene una respuesta estructurada como JSON de OpenAI con sistema
    de reintentos mejorado y estrategias robustas de extracción.

    Args:
        system_msg: Mensaje del sistema para el prompt
        user_msg: Mensaje del usuario para el prompt
        model: Modelo de OpenAI a utilizar
        max_retries: Número máximo de reintentos

    Returns:
        tuple: (contenido raw original, contenido JSON parseado)
    """
    client = get_openai_client()
    if client is None:
        return None, {"error": "Cliente OpenAI no disponible"}

    if not circuit_breaker.can_execute("openai"):
        return None, {"error": "Servicio OpenAI temporalmente no disponible"}

    messages = [
        {"role": "system", "content": system_msg},
        {"role": "user", "content": user_msg}
    ]

    def send_request():
        return client.chat.completions.create(
            model=model,
            temperature=0.5,
            response_format={"type": "json_object"},  # Forzar formato JSON
            messages=messages
        )

    try:
        # Usar retry_with_backoff para gestionar reintentos
        response = retry_with_backoff(send_request, max_retries=max_retries)
        raw_output = response.choices[0].message.content

        # Intentar extraer JSON
        data_json = extract_json_safely(raw_output)

        # Si falló la extracción pero podemos reintentar con un mensaje específico
        if "error" in data_json and max_retries > 0:
            # Añadir mensaje solicitando formato JSON específico
            messages.append({
                "role": "user",
                "content": (
                    "Tu respuesta anterior no cumplió el formato JSON requerido. "
                    "Por favor, responde ÚNICAMENTE en JSON válido con la estructura solicitada. "
                    "No incluyas texto extra, backticks, ni marcadores de código fuente."
                )
            })

            # Reintento específico para corrección de formato
            response = retry_with_backoff(
                lambda: client.chat.completions.create(
                    model=model,
                    temperature=0.3,  # Temperatura más baja para formato más preciso
                    response_format={"type": "json_object"},
                    messages=messages
                ),
                max_retries=1
            )

            # Nuevo intento de extracción
            raw_output = response.choices[0].message.content
            data_json = extract_json_safely(raw_output)

        # Marcar como éxito la comunicación con OpenAI
        circuit_breaker.record_success("openai")
        return raw_output, data_json

    except Exception as e:
        logger.error(f"Error en API de OpenAI: {str(e)}")
        circuit_breaker.record_failure("openai")
        return None, {"error": f"Error en API de OpenAI: {str(e)}"}

# --- 2. INTEGRACIÓN CON ELEVENLABS ---


def generar_audio_consejo(consejo_texto):
    """
    Genera un archivo de audio a partir del texto usando ElevenLabs.

    Args:
        consejo_texto: Texto a convertir en audio

    Returns:
        BytesIO: Buffer con el audio generado, o None si ocurre un error
    """
    if not api_keys["elevenlabs"]["api_key"] or not api_keys["elevenlabs"]["voice_id"]:
        logger.warning("Claves de ElevenLabs no configuradas")
        return None

    if not circuit_breaker.can_execute("elevenlabs"):
        logger.warning("ElevenLabs temporalmente no disponible")
        return None

    if not consejo_texto:
        return None

    # Limpiar el texto
    audio_text = consejo_texto.replace("Consejo final:", "").strip()
    if not audio_text:
        return None

    try:
        elevenlabs_api_key = api_keys["elevenlabs"]["api_key"]
        elevenlabs_voice_id = api_keys["elevenlabs"]["voice_id"]

        tts_url = f"https://api.elevenlabs.io/v1/text-to-speech/{elevenlabs_voice_id}"
        headers = {
            "xi-api-key": elevenlabs_api_key,
            "Content-Type": "application/json"
        }
        data = {
            "text": audio_text,
            "model_id": "eleven_multilingual_v2",
            "voice_settings": {
                "stability": 0.3,
                "similarity_boost": 0.9
            }
        }

        # Función para envío de solicitud
        def send_request():
            response = requests.post(
                tts_url, headers=headers, json=data, timeout=15)
            response.raise_for_status()  # Levantar excepción si hay error
            return response

        # Usar sistema de reintentos
        response_audio = retry_with_backoff(send_request, max_retries=2)

        if response_audio.ok:
            audio_bytes = BytesIO(response_audio.content)
            circuit_breaker.record_success("elevenlabs")
            return audio_bytes
        else:
            logger.error(
                f"Error en ElevenLabs API: {response_audio.status_code}")
            circuit_breaker.record_failure("elevenlabs")
            return None

    except Exception as e:
        logger.error(f"Error al generar audio: {str(e)}")
        circuit_breaker.record_failure("elevenlabs")
        return None

# --- 3. INTEGRACIÓN CON DALL-E ---


def generar_imagen_dalle(tema, nivel):
    """
    Genera una imagen utilizando DALL-E basada en un tema y adaptada al nivel del estudiante.

    Args:
        tema: Tema para la imagen
        nivel: Nivel de español (principiante, intermedio, avanzado)

    Returns:
        tuple: (URL de la imagen generada, descripción de la imagen)
    """
    client = get_openai_client()
    if client is None:
        return None, "API de OpenAI no disponible"

    if not circuit_breaker.can_execute("openai"):
        return None, "Servicio OpenAI temporalmente no disponible"

    # Adaptar la complejidad del prompt según el nivel
    if "principiante" in nivel.lower():
        complejidad = "simple con objetos y personas claramente identificables"
    elif "intermedio" in nivel.lower():
        complejidad = "con detalles moderados y una escena cotidiana con varios elementos"
    else:
        complejidad = "detallada con múltiples elementos, que pueda generar descripciones complejas"

    # Crear el prompt para DALL-E
    prompt = f"Una escena {complejidad} sobre {tema}. La imagen debe ser clara, bien iluminada, y adecuada para describir en español."

    try:
        # Función para envío de solicitud
        def generate_image():
            return client.images.generate(
                model="dall-e-3",
                prompt=prompt,
                n=1,
                size="1024x1024",
                quality="standard"
            )

        # Usar sistema de reintentos
        response = retry_with_backoff(generate_image, max_retries=2)

        # Obtener la URL de la imagen
        imagen_url = response.data[0].url

        # Generar una descripción adaptada al nivel
        descripcion_prompt = f"""
        Crea una descripción en español de esta imagen generada para un estudiante de nivel {nivel}.

        La descripción debe:
        1. Ser apropiada para el nivel {nivel}
        2. Utilizar vocabulario y estructuras gramaticales de ese nivel
        3. Incluir entre 3-5 preguntas al final para que el estudiante practique describiendo la imagen

        Tema de la imagen: {tema}
        """

        def generate_description():
            return client.chat.completions.create(
                model="gpt-4-turbo",
                messages=[
                    {"role": "system", "content": "Eres un profesor de español especializado en crear descripciones y actividades basadas en imágenes."},
                    {"role": "user", "content": descripcion_prompt}
                ],
                temperature=0.7
            )

        descripcion_response = retry_with_backoff(
            generate_description, max_retries=2)
        descripcion = descripcion_response.choices[0].message.content

        # Registrar éxito
        circuit_breaker.record_success("openai")
        return imagen_url, descripcion

    except Exception as e:
        handle_exception("generar_imagen_dalle", e)
        circuit_breaker.record_failure("openai")
        return None, f"Error: {str(e)}"

# --- 4. FUNCIÓN DE OCR PARA TEXTOS MANUSCRITOS ---


def transcribir_imagen_texto(imagen_bytes, idioma="es"):
    """
    Transcribe texto manuscrito de una imagen utilizando la API de OpenAI.
    Versión mejorada con manejo de errores y circuit breaker.

    Args:
        imagen_bytes: Bytes de la imagen a transcribir
        idioma: Código de idioma (es, en, fr)

    Returns:
        str: Texto transcrito o mensaje de error
    """
    client = get_openai_client()
    if client is None:
        return "Error: API de OpenAI no disponible"

    if not circuit_breaker.can_execute("openai"):
        return "Error: Servicio OpenAI temporalmente no disponible"

    try:
        # Codificar la imagen en base64
        encoded_image = base64.b64encode(imagen_bytes).decode('utf-8')

        # Función para envío de solicitud
        def send_ocr_request():
            return client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": f"Eres un sistema de OCR especializado en transcribir texto manuscrito en {idioma}. Tu tarea es extraer con precisión el texto presente en la imagen."
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Transcribe exactamente el texto manuscrito de esta imagen."},
                            {"type": "image_url", "image_url": {
                                "url": f"data:image/jpeg;base64,{encoded_image}"}}
                        ]
                    }
                ],
                max_tokens=1000
            )

        # Usar sistema de reintentos
        response = retry_with_backoff(send_ocr_request, max_retries=2)

        # Registrar éxito
        circuit_breaker.record_success("openai")

        return response.choices[0].message.content.strip()

    except Exception as e:
        logger.error(f"Error en transcribir_imagen_texto: {str(e)}")
        logger.error(traceback.format_exc())

        circuit_breaker.record_failure("openai")

        return f"Error en la transcripción: {str(e)}"

# --- 5. GUARDADO DE DATOS EN GOOGLE SHEETS ---


def guardar_correccion(nombre, nivel, idioma, texto, resultado_json):
    """
    Guarda los datos de una corrección en Google Sheets.

    Args:
        nombre: Nombre del estudiante
        nivel: Nivel de español
        idioma: Idioma de corrección
        texto: Texto original
        resultado_json: Resultado de la corrección en formato JSON (como string o dict)

    Returns:
        dict: Resultado de la operación
    """
    if sheets_connection is None:
        return {"success": False, "message": "Conexión a Google Sheets no disponible"}

    # Fecha actual para el registro
    fecha = datetime.now().strftime("%Y-%m-%d %H:%M")

    # Convertir resultado_json a string si es un diccionario
    if isinstance(resultado_json, dict):
        raw_output = json.dumps(resultado_json)
    else:
        raw_output = resultado_json

    result = {
        "success": True,
        "corrections_saved": False,
        "tracking_saved": False,
        "message": ""
    }

    # Guardar en Historial_Correcciones_ELE
    if sheets_connection["corrections"] is not None:
        try:
            sheets_connection["corrections"].append_row(
                [nombre, nivel, idioma, fecha, texto, raw_output]
            )
            result["corrections_saved"] = True
            logger.info(f"Corrección guardada para {nombre}")
        except Exception as e:
            logger.error(
                f"Error al guardar en Historial_Correcciones_ELE: {str(e)}")
            result["message"] += f"Error al guardar historial: {str(e)}. "

    # Guardar en hoja de seguimiento si hay datos de análisis contextual
    if sheets_connection["tracking"] is not None:
        try:
            # Extraer datos del resultado
            if isinstance(resultado_json, str):
                data_json = extract_json_safely(resultado_json)
            else:
                data_json = resultado_json

            # Extraer datos relevantes para seguimiento
            errores = data_json.get("errores", {})
            analisis = data_json.get("analisis_contextual", {})

            # Contar errores por categoría
            num_gramatica = len(errores.get("Gramática", []))
            num_lexico = len(errores.get("Léxico", []))
            num_puntuacion = len(errores.get("Puntuación", []))
            num_estructura = len(errores.get("Estructura textual", []))
            total_errores = num_gramatica + num_lexico + num_puntuacion + num_estructura

            # Extraer puntuaciones
            coherencia = analisis.get("coherencia", {})
            cohesion = analisis.get("cohesion", {})
            registro = analisis.get("registro_linguistico", {})
            adecuacion = analisis.get("adecuacion_cultural", {})

            puntuacion_coherencia = coherencia.get("puntuacion", 0)
            puntuacion_cohesion = cohesion.get("puntuacion", 0)
            puntuacion_registro = registro.get("puntuacion", 0)
            puntuacion_adecuacion = adecuacion.get("puntuacion", 0)

            # Extraer consejo final
            consejo_final = data_json.get("consejo_final", "")

            # Datos para guardar
            datos_seguimiento = [
                nombre,
                nivel,
                fecha,
                num_gramatica,
                num_lexico,
                num_puntuacion,
                num_estructura,
                total_errores,
                puntuacion_coherencia,
                puntuacion_cohesion,
                puntuacion_registro,
                puntuacion_adecuacion,
                consejo_final
            ]

            # Guardar en hoja de seguimiento
            sheets_connection["tracking"].append_row(datos_seguimiento)
            result["tracking_saved"] = True
            logger.info(f"Estadísticas guardadas para {nombre}")
        except Exception as e:
            logger.error(f"Error al guardar en hoja de Seguimiento: {str(e)}")
            result["message"] += f"Error al guardar estadísticas: {str(e)}."

    # Resultado final
    if result["corrections_saved"] or result["tracking_saved"]:
        result["success"] = True
        if not result["message"]:
            result["message"] = "Datos guardados correctamente."
    else:
        result["success"] = False
        if not result["message"]:
            result["message"] = "No se pudo guardar ningún dato."

    return result


def obtener_historial_estudiante(nombre):
    """
    Obtiene el historial de correcciones y seguimiento para un estudiante específico.

    Args:
        nombre: Nombre del estudiante

    Returns:
        pd.DataFrame o None: DataFrame con historial o None si no hay datos
    """
    if sheets_connection is None or sheets_connection["tracking"] is None:
        logger.warning("No hay conexión con hoja de seguimiento")
        return None

    try:
        # Obtener todos los datos
        todos_datos = sheets_connection["tracking"].get_all_records()

        if not todos_datos:
            return None

        # Crear una versión limpia del nombre buscado
        nombre_buscar = nombre.strip().lower()

        # Buscar en todos los registros con un enfoque más flexible
        datos_estudiante = []
        for row in todos_datos:
            for key, value in row.items():
                # Buscar en cualquier columna que tenga 'nombre'
                if 'nombre' in key.lower() and value:
                    if str(value).strip().lower() == nombre_buscar:
                        datos_estudiante.append(row)
                        break

        # Convertir a DataFrame
        if datos_estudiante:
            df = pd.DataFrame(datos_estudiante)

            # Convertir columnas numéricas explícitamente para evitar errores con PyArrow
            columnas_numericas = [
                'Errores Gramática', 'Errores Léxico', 'Errores Puntuación',
                'Errores Estructura', 'Total Errores', 'Puntuación Coherencia',
                'Puntuación Cohesión', 'Puntuación Registro', 'Puntuación Adecuación Cultural'
            ]

            for col in columnas_numericas:
                if col in df.columns:
                    # Convertir a float de manera segura
                    df[col] = pd.to_numeric(
                        df[col], errors='coerce').fillna(0).astype(float)

            return df

        return None
    except Exception as e:
        logger.error(f"Error en obtener_historial_estudiante: {str(e)}")
        return None
    """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 5: Funciones Utilitarias (no dependientes de UI)
==================================================================================

Este artefacto contiene:
1. Funciones de procesamiento de textos
2. Análisis de complejidad textual
3. Funciones de corrección de texto
4. Generación de consignas y tareas
5. Funciones para exámenes
"""

# --- 1. FUNCIONES DE PROCESAMIENTO DE TEXTOS ---


def generar_consigna_escritura(nivel_actual, tipo_consigna):
    """
    Genera una consigna de escritura adaptada al nivel del estudiante
    y el tipo de texto solicitado.

    Args:
        nivel_actual: Nivel del estudiante (principiante, intermedio, avanzado)
        tipo_consigna: Tipo de texto a generar

    Returns:
        str: Consigna de escritura generada o mensaje de error
    """
    client = get_openai_client()
    if client is None:
        return "No se pudo generar la consigna debido a problemas de conexión."

    if not circuit_breaker.can_execute("openai"):
        return "Servicio temporalmente no disponible. Inténtelo más tarde."

    # Si es aleatorio, seleccionar un tipo
    if tipo_consigna == "Cualquiera (aleatorio)":
        import random
        tipos_disponibles = [
            "Narración", "Correo/Carta formal", "Opinión/Argumentación",
            "Descripción", "Diálogo"
        ]
        tipo_consigna = random.choice(tipos_disponibles)

    # Construir prompt mejorado para OpenAI
    prompt_consigna = f"""
    Eres un profesor experto en la enseñanza de español como lengua extranjera.
    Crea una consigna de escritura adaptada al nivel {nivel_actual} para el tipo de texto: {tipo_consigna}.

    Tu respuesta debe tener este formato exacto:
    1. Un título atractivo y claro
    2. Instrucciones precisas que incluyan:
       - Situación o contexto
       - Tarea específica a realizar
       - Extensión requerida (número de palabras apropiado para el nivel)
       - Elementos que debe incluir el texto

    Adapta la complejidad lingüística y temática al nivel {nivel_actual}:
    - Para niveles principiante: usa vocabulario básico, estructuras simples y temas cotidianos
    - Para niveles intermedio: incluye vocabulario más variado, conectores y temas que requieran opinión
    - Para niveles avanzado: incorpora elementos para expresar matices, argumentación compleja y temas abstractos

    Proporciona solo la consigna, sin explicaciones adicionales ni metacomentarios.
    """

    try:
        def send_request():
            return client.chat.completions.create(
                model="gpt-4-turbo",
                temperature=0.8,
                messages=[
                    {"role": "system", "content": "Eres un profesor de español experto en diseñar actividades de escritura."},
                    {"role": "user", "content": prompt_consigna}
                ]
            )

        # Usar sistema de reintentos
        response = retry_with_backoff(send_request, max_retries=2)

        # Registrar éxito
        circuit_breaker.record_success("openai")
        return response.choices[0].message.content.strip()

    except Exception as e:
        error_msg = f"Error al generar consigna: {str(e)}"
        logger.error(error_msg)
        circuit_breaker.record_failure("openai")
        return f"No se pudo generar la consigna. Error: {str(e)}"


def obtener_criterios_evaluacion(tipo_examen, nivel_examen):
    """
    Obtiene los criterios de evaluación para un examen y nivel específicos.

    Args:
        tipo_examen: Tipo de examen (DELE, SIELE, etc.)
        nivel_examen: Nivel del examen (A1, A2, etc.)

    Returns:
        str: Criterios de evaluación en formato markdown
    """
    # Criterios genéricos por defecto
    criterios_default = """
    ## Criterios de evaluación genéricos

    ### Adecuación al contexto
    - Ajuste a la tarea solicitada
    - Adecuación al registro requerido
    - Cumplimiento del propósito comunicativo

    ### Coherencia y cohesión
    - Organización lógica de las ideas
    - Uso adecuado de conectores
    - Desarrollo temático apropiado

    ### Corrección gramatical
    - Uso adecuado de estructuras gramaticales
    - Control de tiempos verbales
    - Concordancia nominal y verbal

    ### Riqueza léxica
    - Variedad y precisión del vocabulario
    - Uso apropiado de expresiones idiomáticas
    - Evitar repeticiones innecesarias
    """

    # Criterios específicos para DELE
    if tipo_examen == "DELE":
        if nivel_examen in ["A1", "A2"]:
            return """
            ## Criterios de evaluación DELE A1-A2

            ### Adecuación al contexto (25%)
            - Cumple con la tarea solicitada
            - Se ajusta a la extensión requerida
            - Emplea el registro adecuado (formal/informal)

            ### Coherencia textual (25%)
            - Las ideas están organizadas con lógica
            - Usa conectores básicos (y, pero, porque)
            - Información relevante y comprensible

            ### Corrección gramatical (25%)
            - Uso correcto de estructuras básicas
            - Control de presente y pasados simples
            - Concordancia nominal y verbal básica

            ### Alcance y control léxico (25%)
            - Vocabulario básico suficiente
            - Ortografía de palabras frecuentes
            - Expresiones memorizadas adecuadas
            """
        elif nivel_examen in ["B1", "B2"]:
            return """
            ## Criterios de evaluación DELE B1-B2

            ### Adecuación a la tarea (20%)
            - Cumple los puntos requeridos en la tarea
            - Se ajusta a la extensión y formato
            - Registro adecuado al destinatario y propósito

            ### Coherencia y cohesión (20%)
            - Progresión temática clara
            - Uso variado de conectores y marcadores
            - Estructura textual apropiada al género

            ### Corrección gramatical (30%)
            - Estructuras variadas con pocos errores
            - Buen control de tiempos y modos verbales
            - Uso adecuado de subordinación

            ### Alcance y control léxico (30%)
            - Vocabulario preciso y variado
            - Pocas confusiones o imprecisiones léxicas
            - Ortografía y puntuación generalmente correctas
            """
        else:  # C1-C2
            return """
            ## Criterios de evaluación DELE C1-C2

            ### Adecuación a la tarea (20%)
            - Desarrollo completo y matizado de todos los puntos
            - Formato y extensión perfectamente ajustados
            - Registro sofisticado y perfectamente adaptado

            ### Coherencia y cohesión (20%)
            - Estructura textual compleja y elaborada
            - Amplia variedad de mecanismos de cohesión
            - Desarrollo argumentativo sofisticado

            ### Corrección gramatical (30%)
            - Uso preciso y flexible de estructuras complejas
            - Control de aspectos sutiles de la gramática
            - Errores escasos y poco significativos

            ### Alcance y control léxico (30%)
            - Vocabulario amplio, preciso y natural
            - Uso adecuado de expresiones idiomáticas
            - Pleno control de matices y connotaciones
            """

    # Criterios específicos para SIELE
    elif tipo_examen == "SIELE":
        return """
        ## Criterios de evaluación SIELE

        ### Coherencia textual (25%)
        - Organización lógica del contenido
        - Desarrollo adecuado de las ideas
        - Uso de conectores apropiados al nivel

        ### Corrección lingüística (25%)
        - Control gramatical según el nivel
        - Precisión léxica adecuada
        - Ortografía y puntuación

        ### Adecuación al contexto (25%)
        - Cumplimiento de la tarea solicitada
        - Registro apropiado a la situación
        - Longitud del texto según lo requerido

        ### Alcance lingüístico (25%)
        - Variedad de recursos gramaticales
        - Riqueza de vocabulario
        - Complejidad apropiada al nivel
        """

    # Por defecto, devolvemos criterios genéricos
    return criterios_default


def obtener_duracion_examen(tipo_examen, nivel_examen):
    """
    Obtiene la duración en segundos para un simulacro según el tipo y nivel de examen.

    Args:
        tipo_examen: Tipo de examen (DELE, SIELE, etc.)
        nivel_examen: Nivel del examen (A1, A2, etc.)

    Returns:
        int: Duración en segundos
    """
    try:
        # Verificar que los parámetros son válidos
        if tipo_examen is None or nivel_examen is None:
            logger.warning(
                "Parámetros inválidos en obtener_duracion_examen: tipo_examen o nivel_examen es None")
            return 45 * 60  # 45 minutos por defecto

        # Mapeo de duraciones según examen y nivel
        duraciones = {
            "DELE": {
                "A1": 25 * 60,  # 25 minutos
                "A2": 30 * 60,
                "B1": 40 * 60,
                "B2": 60 * 60,
                "C1": 80 * 60,
                "C2": 90 * 60
            },
            "SIELE": {
                "A1": 20 * 60,
                "A2": 25 * 60,
                "B1": 35 * 60,
                "B2": 50 * 60,
                "C1": 70 * 60,
                "C2": 80 * 60
            },
            "CELU": {
                "A1": 30 * 60,
                "A2": 35 * 60,
                "B1": 45 * 60,
                "B2": 60 * 60,
                "C1": 75 * 60,
                "C2": 90 * 60
            },
            "DUCLE": {
                "A1": 20 * 60,
                "A2": 25 * 60,
                "B1": 30 * 60,
                "B2": 40 * 60,
                "C1": 50 * 60,
                "C2": 60 * 60
            }
        }

        # Intentar obtener la duración específica
        duracion = duraciones.get(tipo_examen, {}).get(nivel_examen)

        # Si no se encuentra, usar el valor por defecto
        if duracion is None:
            logger.info(
                f"No se encontró duración para {tipo_examen} nivel {nivel_examen}, usando valor por defecto")
            duracion = 45 * 60  # 45 minutos por defecto

        return duracion

    except Exception as e:
        logger.error(f"Error en obtener_duracion_examen: {str(e)}")
        return 45 * 60  # 45 minutos por defecto


def extraer_titulo(texto):
    """
    Extrae el título de una sección de texto.

    Args:
        texto: Texto de la sección

    Returns:
        str: Título extraído
    """
    if not texto:
        return "Contenido sin título"

    lineas = texto.strip().split("\n")
    if lineas and lineas[0]:
        # Eliminar caracteres no deseados que podrían aparecer en un título
        titulo = lineas[0].strip().strip('#').strip('-').strip('*').strip()
        return titulo

    return "Contenido sin título"

# --- 2. ANÁLISIS DE COMPLEJIDAD TEXTUAL ---


def analizar_complejidad_texto(texto):
    """
    Analiza la complejidad lingüística de un texto en español.

    Args:
        texto: Texto a analizar

    Returns:
        dict: Análisis de complejidad o mensaje de error
    """
    client = get_openai_client()
    if client is None or not texto:
        return {"error": "No se pudo realizar el análisis. Verifique la conexión o el texto."}

    if not circuit_breaker.can_execute("openai"):
        return {"error": "Servicio temporalmente no disponible. Inténtelo más tarde."}

    try:
        # Prompt para análisis de complejidad
        prompt_analisis = f"""
        Analiza la complejidad lingüística del siguiente texto en español.
        Proporciona un análisis detallado que incluya:

        1. Complejidad léxica (variedad de vocabulario, riqueza léxica, palabras poco comunes)
        2. Complejidad sintáctica (longitud de frases, subordinación, tipos de oraciones)
        3. Complejidad textual (coherencia, cohesión, estructura general)
        4. Nivel MCER estimado (A1-C2) con explicación
        5. Índices estadísticos: TTR (type-token ratio), densidad léxica, índice Flesh-Szigriszt (adaptado al español)

        Texto a analizar:
        "{texto}"

        Devuelve el análisis ÚNICAMENTE en formato JSON con la siguiente estructura:
        {{
          "complejidad_lexica": {{
            "nivel": "string",
            "descripcion": "string",
            "palabras_destacadas": ["string1", "string2"]
          }},
          "complejidad_sintactica": {{
            "nivel": "string",
            "descripcion": "string",
            "estructuras_destacadas": ["string1", "string2"]
          }},
          "complejidad_textual": {{
            "nivel": "string",
            "descripcion": "string"
          }},
          "nivel_mcer": {{
            "nivel": "string",
            "justificacion": "string"
          }},
          "indices": {{
            "ttr": number,
            "densidad_lexica": number,
            "szigriszt": number,
            "interpretacion": "string"
          }},
          "recomendaciones": ["string1", "string2"]
        }}
        """

        def send_request():
            return client.chat.completions.create(
                model="gpt-4-turbo",
                temperature=0.3,  # Temperatura baja para resultados más precisos
                # Forzar respuesta en JSON
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": "Eres un experto lingüista y analista textual especializado en complejidad lingüística."},
                    {"role": "user", "content": prompt_analisis}
                ]
            )

        # Usar sistema de reintentos
        response = retry_with_backoff(send_request, max_retries=2)
        raw_output = response.choices[0].message.content

        # Extraer JSON
        analisis_data = extract_json_safely(raw_output)

        # Verificar si se obtuvo un resultado válido
        if "error" in analisis_data:
            circuit_breaker.record_failure("openai")
            return {"error": "No se pudo procesar el análisis. Formato de respuesta incorrecto."}

        # Registrar éxito
        circuit_breaker.record_success("openai")
        return analisis_data

    except Exception as e:
        handle_exception("analizar_complejidad_texto", e)
        circuit_breaker.record_failure("openai")
        return {"error": f"Error al analizar complejidad: {str(e)}"}

# --- 3. FUNCIONES DE CORRECCIÓN DE TEXTO ---


def corregir_texto(texto, nombre, nivel, idioma, tipo_texto, contexto_cultural, info_adicional=""):
    """
    Realiza una corrección completa de un texto con análisis contextual.

    Args:
        texto: Texto a corregir
        nombre: Nombre del estudiante
        nivel: Nivel del estudiante
        idioma: Idioma de corrección (Español, Francés, Inglés)
        tipo_texto: Tipo de texto
        contexto_cultural: Contexto cultural relevante
        info_adicional: Información adicional o contexto

    Returns:
        dict: Resultado de la corrección o mensaje de error
    """
    try:
        client = get_openai_client()
        if client is None:
            return {"error": "Servicio de corrección no disponible. Verifique la conexión."}

        if not circuit_breaker.can_execute("openai"):
            return {"error": "Servicio temporalmente no disponible. Inténtelo más tarde."}

        # Validar la entrada
        if not texto or not nombre:
            return {"error": "El texto y el nombre son obligatorios."}

        # Mapeo de niveles para instrucciones más específicas
        nivel_map_instrucciones = {
            "Nivel principiante (A1-A2)": {
                "descripcion": "principiante (A1-A2)",
                "enfoque": "Enfócate en estructuras básicas, vocabulario fundamental y errores comunes. Utiliza explicaciones simples y claras. Evita terminología lingüística compleja."
            },
            "Nivel intermedio (B1-B2)": {
                "descripcion": "intermedio (B1-B2)",
                "enfoque": "Puedes señalar errores más sutiles de concordancia, uso de tiempos verbales y preposiciones. Puedes usar alguna terminología lingüística básica en las explicaciones."
            },
            "Nivel avanzado (C1-C2)": {
                "descripcion": "avanzado (C1-C2)",
                "enfoque": "Céntrate en matices, coloquialismos, registro lingüístico y fluidez. Puedes usar terminología lingüística específica y dar explicaciones más detalladas y técnicas."
            }
        }

        # Usar nivel intermedio como fallback
        nivel_info = nivel_map_instrucciones.get(
            nivel, nivel_map_instrucciones["Nivel intermedio (B1-B2)"])

        # Instrucciones para el modelo de IA con análisis contextual avanzado
        system_message = f"""
Eres Diego, un profesor experto en ELE (Español como Lengua Extranjera) especializado en análisis lingüístico contextual.
Tu objetivo es corregir textos adaptando tu feedback al nivel {nivel_info['descripcion']} del estudiante.
{nivel_info['enfoque']}

Cuando corrijas un texto, DEBES devolver la respuesta únicamente en un JSON válido, sin texto adicional, con la siguiente estructura EXACTA:

{{
  "saludo": "string",                // en {idioma}
  "tipo_texto": "string",            // en {idioma}
  "errores": {{
       "Gramática": [
           {{
             "fragmento_erroneo": "string",
             "correccion": "string",
             "explicacion": "string"
           }}
           // más errores de Gramática (o [] si ninguno)
       ],
       "Léxico": [
           {{
             "fragmento_erroneo": "string",
             "correccion": "string",
             "explicacion": "string"
           }}
       ],
       "Puntuación": [
           {{
             "fragmento_erroneo": "string",
             "correccion": "string",
             "explicacion": "string"
           }}
       ],
       "Estructura textual": [
           {{
             "fragmento_erroneo": "string",
             "correccion": "string",
             "explicacion": "string"
           }}
       ]
  }},
  "texto_corregido": "string",       // siempre en español
  "analisis_contextual": {{
       "coherencia": {{
           "puntuacion": number,     // del 1 al 10
           "comentario": "string",   // en {idioma}
           "sugerencias": [          // listado de sugerencias en {idioma}
               "string",
               "string"
           ]
       }},
       "cohesion": {{
           "puntuacion": number,     // del 1 al 10
           "comentario": "string",   // en {idioma}
           "sugerencias": [          // listado de sugerencias en {idioma}
               "string",
               "string"
           ]
       }},
       "registro_linguistico": {{
           "puntuacion": number,     // del 1 al 10
           "tipo_detectado": "string", // tipo de registro detectado en {idioma}
           "adecuacion": "string",   // evaluación de adecuación en {idioma}
           "sugerencias": [          // listado de sugerencias en {idioma}
               "string",
               "string"
           ]
       }},
       "adecuacion_cultural": {{
           "puntuacion": number,     // del 1 al 10
           "comentario": "string",   // en {idioma}
           "elementos_destacables": [  // elementos culturales destacables en {idioma}
               "string",
               "string"
           ],
           "sugerencias": [          // listado de sugerencias en {idioma}
               "string",
               "string"
           ]
       }}
  }},
  "consejo_final": "string",         // en español
  "fin": "Fin de texto corregido."
}}

IMPORTANTE:
- Las explicaciones de los errores deben estar en {idioma}
- Todo el análisis contextual debe estar en {idioma}
- El texto corregido completo SIEMPRE debe estar en español, independientemente del idioma seleccionado
- El consejo final SIEMPRE debe estar en español
- Adapta tus explicaciones y sugerencias al nivel {nivel_info['descripcion']} del estudiante
- Considera el tipo de texto "{tipo_texto}" y el contexto cultural "{contexto_cultural}" en tu análisis

No devuelvas ningún texto extra fuera de este JSON.
"""

        # Mensaje para el usuario con contexto adicional
        user_message = f"""
Texto del alumno:
\"\"\"
{texto}
\"\"\"
Nivel: {nivel}
Nombre del alumno: {nombre}
Idioma de corrección: {idioma}
Tipo de texto: {tipo_texto}
Contexto cultural: {contexto_cultural}
{f"Información adicional: {info_adicional}" if info_adicional else ""}
"""

        try:
            # Enviar solicitud a OpenAI
            raw_output, data_json = obtener_json_de_ia(
                system_message, user_message, model="gpt-4-turbo", max_retries=3)

            # Verificar si hay error en la respuesta
            if raw_output is None or "error" in data_json:
                error_msg = data_json.get(
                    "error", "Error desconocido en el procesamiento")
                logger.error(f"Error en corrección: {error_msg}")
                return {"error": error_msg}

            # Registrar éxito
            circuit_breaker.record_success("openai")

            # Guardar corrección si hay conexión a Google Sheets
            if sheets_connection is not None:
                resultado_guardado = guardar_correccion(
                    nombre, nivel, idioma, texto, raw_output)
                if not resultado_guardado["success"]:
                    logger.warning(
                        f"No se pudo guardar la corrección: {resultado_guardado['message']}")

            # Guardar el texto para posible uso futuro
            set_session_var("ultimo_texto", texto)

            # Devolver resultado
            return data_json

        except Exception as e:
            error_msg = f"Error al corregir texto: {str(e)}"
            logger.error(error_msg)
            circuit_breaker.record_failure("openai")
            return {"error": error_msg}

    except Exception as e:
        handle_exception("corregir_texto", e)
        return {"error": f"Error al corregir texto: {str(e)}"}


def corregir_examen(texto, tipo_examen, nivel_examen, tiempo_usado=None):
    """
    Corrige un texto de examen específico.

    Args:
        texto: Texto a corregir
        tipo_examen: Tipo de examen
        nivel_examen: Nivel del examen
        tiempo_usado: Tiempo usado (opcional)

    Returns:
        dict: Resultado de la corrección
    """
    if not texto or not texto.strip():
        return {"error": "El texto está vacío. Por favor, escribe una respuesta."}

    # Guardar para futura referencia
    set_session_var("ultimo_texto", texto)

    # Obtener datos necesarios
    nombre = get_session_var("usuario_actual", "Usuario")

    # Mapear nivel del examen al formato de nivel de la aplicación
    nivel_map = {
        "A1": "Nivel principiante (A1-A2)",
        "A2": "Nivel principiante (A1-A2)",
        "B1": "Nivel intermedio (B1-B2)",
        "B2": "Nivel intermedio (B1-B2)",
        "C1": "Nivel avanzado (C1-C2)",
        "C2": "Nivel avanzado (C1-C2)"
    }
    nivel = nivel_map.get(nivel_examen, "Nivel intermedio (B1-B2)")

    # Construir información adicional
    info_adicional = f"Examen {tipo_examen} nivel {nivel_examen}"
    if tiempo_usado:
        info_adicional += f" (Tiempo usado: {tiempo_usado})"

    # Llamar a la función de corrección
    resultado = corregir_texto(
        texto, nombre, nivel, "Español", "Académico",
        "Contexto académico", info_adicional
    )

    # Guardar resultado
    set_session_var("correction_result", resultado)
    set_session_var("last_correction_time", datetime.now().isoformat())
    set_session_var("examen_result", resultado)

    return resultado


def corregir_descripcion_imagen(descripcion, tema_imagen, nivel):
    """
    Corrige una descripción de imagen.

    Args:
        descripcion: Texto de la descripción
        tema_imagen: Tema de la imagen
        nivel: Nivel del estudiante

    Returns:
        dict: Resultado de la corrección
    """
    if not descripcion or not descripcion.strip():
        return {"error": "La descripción está vacía. Por favor, escribe una descripción."}

    # Guardar para futura referencia
    set_session_var("ultimo_texto", descripcion)

    # Obtener datos necesarios
    nombre = get_session_var("usuario_actual", "Usuario")

    # Información adicional
    info_adicional = f"Descripción de imagen sobre '{tema_imagen}'"

    # Llamar a la función de corrección
    resultado = corregir_texto(
        descripcion, nombre, nivel, "Español", "Descriptivo",
        "General/Internacional", info_adicional
    )

    # Guardar resultado
    set_session_var("correction_result", resultado)
    set_session_var("last_correction_time", datetime.now().isoformat())

    return resultado


def generar_tarea_examen(tipo_examen, nivel_examen):
    """
    Genera una tarea de expresión escrita para un examen específico.

    Args:
        tipo_examen: Tipo de examen (DELE, SIELE, etc.)
        nivel_examen: Nivel del examen (A1, A2, etc.)

    Returns:
        str: Tarea generada para el examen
    """
    client = get_openai_client()
    if client is None:
        return "No se pudo generar la tarea. Servicio no disponible."

    if not circuit_breaker.can_execute("openai"):
        return "Servicio temporalmente no disponible. Inténtelo más tarde."

    try:
        # Prompt para generación de tarea
        prompt_tarea = f"""
        Crea una tarea de expresión escrita para el examen {tipo_examen} de nivel {nivel_examen}.
        La tarea debe incluir:
        1. Instrucciones claras y precisas
        2. Contexto o situación comunicativa
        3. Número de palabras requerido
        4. Aspectos que se evaluarán

        El formato debe ser idéntico al que aparece en los exámenes oficiales {tipo_examen}.
        La tarea debe ser apropiada para el nivel {nivel_examen}, siguiendo los estándares oficiales.
        """

        def send_request():
            return client.chat.completions.create(
                model="gpt-4-turbo",
                temperature=0.7,
                messages=[
                    {"role": "system", "content": "Eres un experto en exámenes oficiales de español como lengua extranjera."},
                    {"role": "user", "content": prompt_tarea}
                ]
            )

        # Usar sistema de reintentos
        response = retry_with_backoff(send_request, max_retries=2)

        # Registrar éxito
        circuit_breaker.record_success("openai")
        return response.choices[0].message.content

    except Exception as e:
        error_msg = f"Error al generar tarea de examen: {str(e)}"
        logger.error(error_msg)
        circuit_breaker.record_failure("openai")
        return f"No se pudo generar la tarea. Error: {str(e)}"


def generar_ejemplos_evaluados(tipo_examen, nivel_examen):
    """
    Genera ejemplos de textos evaluados para un examen específico.

    Args:
        tipo_examen: Tipo de examen (DELE, SIELE, etc.)
        nivel_examen: Nivel del examen (A1, A2, etc.)

    Returns:
        str: Ejemplos de textos evaluados
    """
    client = get_openai_client()
    if client is None:
        return "No se pudieron generar ejemplos. Servicio no disponible."

    if not circuit_breaker.can_execute("openai"):
        return "Servicio temporalmente no disponible. Inténtelo más tarde."

    try:
        # Prompt para generación de ejemplos
        prompt_ejemplos = f"""
        Genera un ejemplo de texto de un estudiante para el examen {tipo_examen} nivel {nivel_examen},
        junto con una evaluación detallada usando los criterios oficiales.
        Muestra:
        1. La tarea solicitada
        2. El texto del estudiante (con algunos errores típicos de ese nivel)
        3. Evaluación punto por punto según los criterios oficiales
        4. Puntuación desglosada y comentarios
        """

        def send_request():
            return client.chat.completions.create(
                model="gpt-4-turbo",
                temperature=0.7,
                messages=[
                    {"role": "system", "content": "Eres un evaluador experto de exámenes oficiales de español."},
                    {"role": "user", "content": prompt_ejemplos}
                ]
            )

        # Usar sistema de reintentos
        response = retry_with_backoff(send_request, max_retries=2)

        # Registrar éxito
        circuit_breaker.record_success("openai")
        return response.choices[0].message.content

    except Exception as e:
        error_msg = f"Error al generar ejemplos evaluados: {str(e)}"
        logger.error(error_msg)
        circuit_breaker.record_failure("openai")
        return f"No se pudieron generar ejemplos. Error: {str(e)}"
    """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 6: Componentes Reutilizables de UI Básicos
==================================================================================

Este artefacto contiene:
1. Componentes básicos de UI reutilizables
2. Utilidades de interfaz
3. Mensajes estandarizados
4. Indicadores de progreso
5. Diálogos de confirmación
"""

# --- 1. COMPONENTES DE UI REUTILIZABLES ---


def ui_header():
    """Muestra el encabezado principal de la aplicación."""
    col1, col2 = st.columns([6, 1])
    with col1:
        st.title("📝 Textocorrector ELE")
        st.markdown(
            "Corrección de textos en español con análisis contextual avanzado.")

    with col2:
        # Mostrar indicador de versión
        st.markdown(f"""
        <div style="background-color:#f0f2f6;padding:8px;border-radius:5px;margin-top:20px;text-align:center">
            <small>v{APP_VERSION}</small>
        </div>
        """, unsafe_allow_html=True)


def ui_user_info_form(form_key="form_user_info"):
    """
    Formulario para obtener información básica del usuario.

    Args:
        form_key: Clave única para el formulario

    Returns:
        dict: Datos del usuario (nombre, nivel)
    """
    with st.form(key=form_key):
        col1, col2 = st.columns(2)

        with col1:
            nombre = st.text_input(
                "Nombre y apellido:",
                value=get_session_var("usuario_actual", ""),
                help="Por favor, introduce tanto tu nombre como tu apellido separados por un espacio."
            )

        with col2:
            nivel = st.selectbox(
                "¿Cuál es tu nivel?",
                [
                    "Nivel principiante (A1-A2)",
                    "Nivel intermedio (B1-B2)",
                    "Nivel avanzado (C1-C2)"
                ],
                index=["principiante", "intermedio", "avanzado"].index(
                    get_session_var("nivel_estudiante", "intermedio")
                )
            )

        submit = st.form_submit_button("Guardar", use_container_width=True)

        if submit:
            # Validar nombre
            if not nombre or " " not in nombre:
                st.warning(
                    "Por favor, introduce tanto el nombre como el apellido separados por un espacio.")
                return None

            # Guardar en session_state
            set_session_var("usuario_actual", nombre)

            # Guardar nivel en formato simplificado
            nivel_map = {
                "Nivel principiante (A1-A2)": "principiante",
                "Nivel intermedio (B1-B2)": "intermedio",
                "Nivel avanzado (C1-C2)": "avanzado"
            }
            set_session_var("nivel_estudiante",
                            nivel_map.get(nivel, "intermedio"))

            return {"nombre": nombre, "nivel": nivel}

        return None


def ui_idioma_correcciones_tipo():
    """
    Componente para seleccionar idioma de correcciones y tipo de texto.

    Returns:
        dict: Opciones seleccionadas
    """
    col1, col2, col3 = st.columns(3)

    with col1:
        idioma = st.selectbox(
            "Idioma de corrección",
            ["Español", "Inglés", "Francés"],
            help="Idioma en el que recibirás las explicaciones y análisis."
        )

    with col2:
        tipo_texto = st.selectbox(
            "Tipo de texto",
            [
                "General/No especificado",
                "Académico",
                "Profesional/Laboral",
                "Informal/Cotidiano",
                "Creativo/Literario"
            ],
            help="Tipo de texto que estás escribiendo."
        )

    with col3:
        contexto_cultural = st.selectbox(
            "Contexto cultural",
            [
                "General/Internacional",
                "España",
                "Latinoamérica",
                "Contexto académico",
                "Contexto empresarial"
            ],
            help="Contexto cultural relevante para tu texto."
        )

    return {
        "idioma": idioma,
        "tipo_texto": tipo_texto,
        "contexto_cultural": contexto_cultural
    }


def ui_examen_options():
    """
    Componente para seleccionar opciones de examen.

    Returns:
        dict: Opciones de examen seleccionadas
    """
    col1, col2 = st.columns(2)

    with col1:
        tipo_examen = st.selectbox(
            "Examen oficial:",
            ["DELE", "SIELE", "CELU", "DUCLE"],
            help="Selecciona el tipo de examen para el que quieres prepararte."
        )

    with col2:
        nivel_examen = st.selectbox(
            "Nivel:",
            ["A1", "A2", "B1", "B2", "C1", "C2"],
            help="Nivel del examen."
        )

    return {
        "tipo_examen": tipo_examen,
        "nivel_examen": nivel_examen
    }


def ui_loading_spinner(text="Procesando..."):
    """
    Spinner de carga con texto personalizable.

    Args:
        text: Texto a mostrar durante la carga

    Returns:
        st.spinner: Objeto spinner de Streamlit
    """
    return st.spinner(text)


def ui_empty_placeholder():
    """
    Crea un placeholder vacío para contenido dinámico.

    Returns:
        st.empty: Objeto empty de Streamlit
    """
    return st.empty()


def ui_countdown_timer(total_seconds, start_time=None):
    """
    Muestra un temporizador de cuenta regresiva.

    Args:
        total_seconds: Tiempo total en segundos
        start_time: Tiempo de inicio (None = ahora)

    Returns:
        dict: Estado del temporizador
    """
    # Manejar el caso donde total_seconds es None
    if total_seconds is None:
        total_seconds = 0  # Usar 0 como valor por defecto

    if start_time is None:
        start_time = time.time()

    # Calcular tiempo transcurrido
    tiempo_transcurrido = time.time() - start_time
    tiempo_restante_segundos = max(0, total_seconds - tiempo_transcurrido)

    # Formatear tiempo restante
    minutos = int(tiempo_restante_segundos // 60)
    segundos = int(tiempo_restante_segundos % 60)
    tiempo_formateado = f"{minutos:02d}:{segundos:02d}"

    # Calcular porcentaje (evitar división por cero)
    if total_seconds > 0:
        porcentaje = 1 - (tiempo_restante_segundos / total_seconds)
    else:
        porcentaje = 1  # Si no hay tiempo total, consideramos que está completo

    porcentaje = max(0, min(1, porcentaje))  # Asegurar entre 0 y 1

    # Determinar color según tiempo restante
    if tiempo_restante_segundos > total_seconds * 0.5:  # Más del 50% restante
        color = "normal"  # Verde/Normal
    elif tiempo_restante_segundos > total_seconds * 0.25:  # Entre 25% y 50%
        color = "warning"  # Amarillo/Advertencia
    else:  # Menos del 25%
        color = "error"  # Rojo/Error

    return {
        "tiempo_restante": tiempo_restante_segundos,
        "tiempo_formateado": tiempo_formateado,
        "porcentaje": porcentaje,
        "color": color,
        "terminado": tiempo_restante_segundos <= 0
    }

# --- 2. UTILIDADES DE INTERFAZ ---


def ui_error_message(error_msg, show_details=True):
    """
    Muestra un mensaje de error formateado.

    Args:
        error_msg: Mensaje de error
        show_details: Mostrar detalles adicionales
    """
    st.error(f"⚠️ {error_msg}")

    if show_details:
        with st.expander("Ver detalles del error"):
            st.code(traceback.format_exc())
            st.info(
                "Si el problema persiste, contacta con el administrador del sistema.")


def ui_success_message(msg):
    """
    Muestra un mensaje de éxito formateado.

    Args:
        msg: Mensaje de éxito
    """
    st.success(f"✅ {msg}")


def ui_info_message(msg):
    """
    Muestra un mensaje informativo formateado.

    Args:
        msg: Mensaje informativo
    """
    st.info(f"ℹ️ {msg}")


def ui_warning_message(msg):
    """
    Muestra un mensaje de advertencia formateado.

    Args:
        msg: Mensaje de advertencia
    """
    st.warning(f"⚠️ {msg}")


def ui_show_progress(title, value, max_value=100, style="progress"):
    """
    Muestra una barra de progreso con diferentes estilos.

    Args:
        title: Título del progreso
        value: Valor actual
        max_value: Valor máximo
        style: Estilo (progress/metric/percent)
    """
    if style == "progress":
        st.markdown(f"#### {title}")
        st.progress(value / max_value)
    elif style == "metric":
        st.metric(title, f"{value}/{max_value}")
    elif style == "percent":
        percent = (value / max_value) * 100
        st.metric(title, f"{percent:.0f}%")
    else:
        st.markdown(f"**{title}:** {value}/{max_value}")


def ui_confirm_dialog(title, message, ok_button="Confirmar", cancel_button="Cancelar"):
    """
    Muestra un diálogo de confirmación.

    Args:
        title: Título del diálogo
        message: Mensaje del diálogo
        ok_button: Texto del botón de confirmación
        cancel_button: Texto del botón de cancelación

    Returns:
        bool: True si se confirma, False si se cancela
    """
    st.markdown(f"### {title}")
    st.markdown(message)

    col1, col2 = st.columns(2)
    with col1:
        cancel = st.button(cancel_button, key=f"cancel_{hash(title)}")
    with col2:
        confirm = st.button(ok_button, key=f"confirm_{hash(title)}")

    if confirm:
        return True

    if cancel:
        return False

    return None  # No se ha tomado decisión


def ui_tooltip(text, tooltip):
    """
    Muestra un texto con tooltip al pasar el ratón.

    Args:
        text: Texto a mostrar
        tooltip: Texto del tooltip
    """
    st.markdown(f"""
    <span title="{tooltip}" style="border-bottom: 1px dotted #000; cursor: help;">
        {text}
    </span>
    """, unsafe_allow_html=True)


def ui_feedback_form():
    """
    Muestra un formulario de feedback para el usuario.

    Returns:
        dict: Datos del feedback o None si no se envía
    """
    with st.expander("📝 Danos tu opinión", expanded=False):
        with st.form(key="feedback_form"):
            st.markdown("### Nos gustaría conocer tu opinión")

            rating = st.slider(
                "¿Cómo valorarías la utilidad de esta herramienta?",
                min_value=1,
                max_value=5,
                value=4,
                help="1 = Poco útil, 5 = Muy útil"
            )

            feedback_text = st.text_area(
                "Comentarios o sugerencias:",
                height=100,
                help="¿Qué podríamos mejorar?"
            )

            submit = st.form_submit_button("Enviar feedback")

            if submit:
                # En una implementación real, aquí se enviaría el feedback a una base de datos
                ui_success_message("¡Gracias por tu feedback!")
                return {
                    "rating": rating,
                    "feedback": feedback_text,
                    "timestamp": datetime.now().isoformat()
                }

            return None
        """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 7: Funciones de Procesamiento que Dependen de UI
==================================================================================

Este artefacto contiene:
1. Generación de informes en diferentes formatos
2. Visualización de datos y estadísticas
3. Base de datos de recursos educativos
4. Generación de ejercicios personalizados
5. Generación de planes de estudio
"""

# --- 1. GENERACIÓN DE INFORMES EN DIFERENTES FORMATOS ---


def generar_informe_docx(nombre, nivel, fecha, texto_original, texto_corregido, errores_obj, analisis_contextual, consejo_final):
    """
    Genera un informe de corrección en formato Word (DOCX).
    Versión mejorada con mejor manejo de errores y validación.

    Args:
        nombre: Nombre del estudiante
        nivel: Nivel del estudiante
        fecha: Fecha de la corrección
        texto_original: Texto original del estudiante
        texto_corregido: Texto con correcciones
        errores_obj: Objeto con errores detectados
        analisis_contextual: Objeto con análisis contextual
        consejo_final: Consejo final para el estudiante

    Returns:
        BytesIO: Buffer con el documento generado
    """
    try:
        # Añadir información de depuración
        logger.info(f"Iniciando generación de informe DOCX para {nombre}")

        # Crear el documento desde cero
        doc = Document()

        # Estilo del documento
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        # Título
        doc.add_heading('Informe de corrección textual', 0)

        # Información general
        doc.add_heading('Información general', level=1)
        doc.add_paragraph(f'Nombre: {nombre}')
        doc.add_paragraph(f'Nivel: {nivel}')
        doc.add_paragraph(f'Fecha: {fecha}')

        # Texto original
        doc.add_heading('Texto original', level=1)
        doc.add_paragraph(texto_original)

        # Texto corregido
        doc.add_heading('Texto corregido', level=1)
        doc.add_paragraph(texto_corregido)

        # Análisis de errores
        doc.add_heading('Análisis de errores', level=1)

        # Verificar que errores_obj es un diccionario
        if errores_obj and isinstance(errores_obj, dict):
            for categoria, errores in errores_obj.items():
                if errores and isinstance(errores, list) and len(errores) > 0:
                    doc.add_heading(categoria, level=2)
                    for error in errores:
                        if isinstance(error, dict):
                            p = doc.add_paragraph()

                            # Verificar que los campos existan antes de agregarlos
                            fragmento = error.get('fragmento_erroneo', '')
                            if fragmento:
                                run = p.add_run('Fragmento erróneo: ')
                                run.bold = True
                                run = p.add_run(fragmento)
                                run.font.color.rgb = RGBColor(255, 0, 0)

                            correccion = error.get('correccion', '')
                            if correccion:
                                p = doc.add_paragraph()
                                run = p.add_run('Corrección: ')
                                run.bold = True
                                run = p.add_run(correccion)
                                run.font.color.rgb = RGBColor(0, 128, 0)

                            explicacion = error.get('explicacion', '')
                            if explicacion:
                                p = doc.add_paragraph()
                                run = p.add_run('Explicación: ')
                                run.bold = True
                                p.add_run(explicacion)

                            doc.add_paragraph()  # Espacio
        else:
            doc.add_paragraph("No se detectaron errores significativos.")

        # Análisis contextual
        doc.add_heading('Análisis contextual', level=1)

        # Tabla de puntuaciones
        doc.add_heading('Puntuaciones', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Aspecto'
        hdr_cells[1].text = 'Coherencia'
        hdr_cells[2].text = 'Cohesión'
        hdr_cells[3].text = 'Registro'
        hdr_cells[4].text = 'Adecuación cultural'

        # Verificar que analisis_contextual es un diccionario
        if analisis_contextual and isinstance(analisis_contextual, dict):
            # Datos
            row_cells = table.add_row().cells
            row_cells[0].text = 'Puntuación'

            # Obtener puntuaciones con manejo seguro
            coherencia = analisis_contextual.get('coherencia', {})
            cohesion = analisis_contextual.get('cohesion', {})
            registro = analisis_contextual.get('registro_linguistico', {})
            adecuacion = analisis_contextual.get('adecuacion_cultural', {})

            row_cells[1].text = str(coherencia.get('puntuacion', 'N/A'))
            row_cells[2].text = str(cohesion.get('puntuacion', 'N/A'))
            row_cells[3].text = str(registro.get('puntuacion', 'N/A'))
            row_cells[4].text = str(adecuacion.get('puntuacion', 'N/A'))

            # Añadir comentarios del análisis contextual
            if coherencia:
                doc.add_heading('Coherencia textual', level=3)
                doc.add_paragraph(coherencia.get(
                    'comentario', 'No disponible'))

            if cohesion:
                doc.add_heading('Cohesión textual', level=3)
                doc.add_paragraph(cohesion.get('comentario', 'No disponible'))

            if registro:
                doc.add_heading('Registro lingüístico', level=3)
                doc.add_paragraph(
                    f"Tipo detectado: {registro.get('tipo_detectado', 'No especificado')}")
                doc.add_paragraph(registro.get('adecuacion', 'No disponible'))

            if adecuacion:
                doc.add_heading('Adecuación cultural', level=3)
                doc.add_paragraph(adecuacion.get(
                    'comentario', 'No disponible'))
        else:
            doc.add_paragraph("Análisis contextual no disponible.")

        # Consejo final
        doc.add_heading('Consejo final', level=1)
        doc.add_paragraph(consejo_final or "No disponible")

        # CAMBIO: Simplificar generación de QR para evitar problemas
        try:
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )

            # Generar un ID único para el informe
            informe_id = f"{nombre.replace(' ', '')}_{fecha.replace(' ', '_').replace(':', '-')}"
            qr_data = f"textocorrector://informe/{informe_id}"
            qr.add_data(qr_data)
            qr.make(fit=True)

            img = qr.make_image(fill_color="black", back_color="white")

            # Guardar QR como imagen temporal
            qr_buffer = BytesIO()
            img.save(qr_buffer)
            qr_buffer.seek(0)

            # Añadir la imagen del QR al documento
            doc.add_heading('Acceso online', level=1)
            doc.add_paragraph(
                'Escanea este código QR para acceder a este informe online:')

            # CAMBIO: Usar un enfoque más robusto para añadir la imagen
            try:
                doc.add_picture(qr_buffer, width=Inches(2.0))
            except Exception as pic_error:
                logger.error(
                    f"Error al añadir QR como imagen: {str(pic_error)}")
                doc.add_paragraph(
                    "Código QR no disponible - Error al generar imagen")

            # Cerrar el buffer del QR
            qr_buffer.close()
            logger.info("Código QR generado correctamente")
        except Exception as qr_error:
            logger.error(f"Error al generar QR: {str(qr_error)}")
            # Continuar sin el QR
            doc.add_heading('Acceso online', level=1)
            doc.add_paragraph('Código QR no disponible en este momento.')

        # Guardar el documento en memoria
        logger.info("Guardando documento DOCX en memoria")
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Verificar tamaño del buffer
        buffer_size = len(docx_buffer.getvalue())
        logger.info(
            f"Documento DOCX generado correctamente. Tamaño: {buffer_size} bytes")

        if buffer_size == 0:
            logger.error("¡El buffer del documento tiene tamaño cero!")
            raise ValueError("El documento generado está vacío")

        return docx_buffer

    except Exception as e:
        logger.error(f"Error detallado al generar informe DOCX: {str(e)}")
        logger.error(traceback.format_exc())

        # Mostrar error detallado sin interrumpir la aplicación
        try:
            # Crear un documento de error
            error_doc = Document()
            error_doc.add_heading('Error al generar informe', 0)
            error_doc.add_paragraph(
                f"Se produjo un error al generar el informe: {str(e)}")
            error_doc.add_paragraph(
                f"Detalles técnicos: {traceback.format_exc()[:500]}...")
            error_doc.add_paragraph("Por favor, contacte con soporte técnico.")

            # Salvar documento de error
            error_buffer = BytesIO()
            error_doc.save(error_buffer)
            error_buffer.seek(0)
            return error_buffer
        except Exception as inner_e:
            logger.error(
                f"Error secundario al generar informe de error: {str(inner_e)}")
            # Si falla completamente, devolver None
            return None
        """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 7: Funciones de Procesamiento que Dependen de UI (continuación 1)
==================================================================================

Continuación de las funciones de procesamiento que dependen de la UI
"""


def generar_informe_html(nombre, nivel, fecha, texto_original, texto_corregido, analisis_contextual, consejo_final):
    """
    Genera un informe de corrección en formato HTML.
    Versión mejorada con mejor manejo de valores nulos y formato.

    Args:
        nombre: Nombre del estudiante
        nivel: Nivel del estudiante
        fecha: Fecha de la corrección
        texto_original: Texto original del estudiante
        texto_corregido: Texto con correcciones
        analisis_contextual: Objeto con análisis contextual
        consejo_final: Consejo final para el estudiante

    Returns:
        str: Contenido HTML del informe
    """
    try:
        # Verificar entradas con valores seguros por defecto
        nombre = nombre or "Estudiante"
        nivel = nivel or "No especificado"
        fecha = fecha or datetime.now().strftime("%Y-%m-%d %H:%M")
        texto_original = texto_original or "No disponible"
        texto_corregido = texto_corregido or "No disponible"
        consejo_final = consejo_final or "No disponible"
        app_version = APP_VERSION  # Usar variable global

        # Sanitizar textos para HTML
        def sanitize_html(text):
            if not text:
                return ""
            # Reemplazar caracteres problemáticos
            sanitized = text.replace("<", "&lt;").replace(">", "&gt;")
            # Convertir saltos de línea en <br>
            sanitized = sanitized.replace("\n", "<br>")
            return sanitized

        texto_original_safe = sanitize_html(texto_original)
        texto_corregido_safe = sanitize_html(texto_corregido)
        consejo_final_safe = sanitize_html(consejo_final)

        # Verificar analisis_contextual
        if not isinstance(analisis_contextual, dict):
            analisis_contextual = {}

        # Obtener puntuaciones con manejo seguro
        coherencia = analisis_contextual.get('coherencia', {})
        cohesion = analisis_contextual.get('cohesion', {})
        registro = analisis_contextual.get('registro_linguistico', {})
        adecuacion = analisis_contextual.get('adecuacion_cultural', {})

        puntuacion_coherencia = coherencia.get('puntuacion', 'N/A')
        puntuacion_cohesion = cohesion.get('puntuacion', 'N/A')
        puntuacion_registro = registro.get('puntuacion', 'N/A')
        puntuacion_adecuacion = adecuacion.get('puntuacion', 'N/A')

        # Crear HTML con estructura mejorada
        html_content = f'''
        <!DOCTYPE html>
        <html lang="es">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Informe de corrección - {nombre}</title>
            <style>
                body {{ 
                    font-family: Arial, sans-serif; 
                    line-height: 1.6; 
                    margin: 0;
                    padding: 20px;
                    color: #333;
                }}
                .container {{ 
                    max-width: 800px; 
                    margin: 0 auto; 
                    padding: 20px; 
                    box-shadow: 0 0 10px rgba(0,0,0,0.1);
                    background-color: #fff;
                    border-radius: 5px;
                }}
                h1 {{ 
                    color: #2c3e50; 
                    border-bottom: 2px solid #3498db;
                    padding-bottom: 10px;
                }}
                h2 {{ 
                    color: #3498db; 
                    margin-top: 30px; 
                    border-left: 4px solid #3498db;
                    padding-left: 10px;
                }}
                h3 {{ 
                    color: #2980b9; 
                    margin-top: 20px;
                }}
                .original {{ 
                    background-color: #f8f9fa; 
                    padding: 15px; 
                    border-left: 4px solid #6c757d; 
                    white-space: pre-wrap;
                    margin: 15px 0;
                    border-radius: 4px;
                }}
                .corregido {{ 
                    background-color: #e7f4e4; 
                    padding: 15px; 
                    border-left: 4px solid #28a745; 
                    white-space: pre-wrap;
                    margin: 15px 0;
                    border-radius: 4px;
                }}
                .error-item {{ 
                    margin-bottom: 20px; 
                    padding: 10px; 
                    background-color: #f1f1f1;
                    border-radius: 4px;
                }}
                .fragmento {{ 
                    color: #dc3545; 
                    font-weight: bold;
                }}
                .correccion {{ 
                    color: #28a745; 
                    font-weight: bold;
                }}
                .explicacion {{ 
                    color: #17a2b8; 
                    font-style: italic; 
                    margin-top: 10px;
                }}
                .puntuaciones {{ 
                    width: 100%; 
                    border-collapse: collapse; 
                    margin: 20px 0;
                }}
                .puntuaciones th, .puntuaciones td {{ 
                    border: 1px solid #ddd; 
                    padding: 10px; 
                    text-align: center;
                }}
                .puntuaciones th {{ 
                    background-color: #f2f2f2; 
                    font-weight: bold;
                }}
                .consejo {{ 
                    background-color: #e7f5fe; 
                    padding: 15px; 
                    border-left: 4px solid #17a2b8; 
                    margin: 20px 0;
                    border-radius: 4px;
                }}
                .footer {{ 
                    margin-top: 50px; 
                    padding-top: 20px; 
                    border-top: 1px solid #ddd; 
                    color: #6c757d; 
                    font-size: 0.8em;
                    text-align: center;
                }}
                @media print {{
                    body {{ background-color: #fff; }}
                    .container {{ box-shadow: none; }}
                    a {{ text-decoration: none; color: #000; }}
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <h1>Informe de corrección textual</h1>

                <section>
                    <h2>Información general</h2>
                    <p><strong>Nombre:</strong> {nombre}</p>
                    <p><strong>Nivel:</strong> {nivel}</p>
                    <p><strong>Fecha:</strong> {fecha}</p>
                </section>

                <section>
                    <h2>Texto original</h2>
                    <div class="original">
                        {texto_original_safe}
                    </div>

                    <h2>Texto corregido</h2>
                    <div class="corregido">
                        {texto_corregido_safe}
                    </div>
                </section>

                <section>
                    <h2>Análisis contextual</h2>

                    <h3>Puntuaciones</h3>
                    <table class="puntuaciones">
                        <tr>
                            <th>Coherencia</th>
                            <th>Cohesión</th>
                            <th>Registro</th>
                            <th>Adecuación cultural</th>
                        </tr>
                        <tr>
                            <td>{puntuacion_coherencia}/10</td>
                            <td>{puntuacion_cohesion}/10</td>
                            <td>{puntuacion_registro}/10</td>
                            <td>{puntuacion_adecuacion}/10</td>
                        </tr>
                    </table>
                    
                    <h3>Coherencia textual</h3>
                    <p>{sanitize_html(coherencia.get('comentario', 'No disponible'))}</p>
                    
                    <h3>Cohesión textual</h3>
                    <p>{sanitize_html(cohesion.get('comentario', 'No disponible'))}</p>
                    
                    <h3>Registro lingüístico</h3>
                    <p><strong>Tipo detectado:</strong> {sanitize_html(registro.get('tipo_detectado', 'No especificado'))}</p>
                    <p>{sanitize_html(registro.get('adecuacion', 'No disponible'))}</p>
                    
                    <h3>Adecuación cultural</h3>
                    <p>{sanitize_html(adecuacion.get('comentario', 'No disponible'))}</p>
                </section>

                <section>
                    <h2>Consejo final</h2>
                    <div class="consejo">
                        <p>{consejo_final_safe}</p>
                    </div>
                </section>

                <div class="footer">
                    <p>Textocorrector ELE - Informe generado el {fecha} - Versión {app_version}</p>
                    <p>Este informe fue generado automáticamente por la aplicación Textocorrector ELE</p>
                </div>
            </div>
        </body>
        </html>
        '''

        return html_content

    except Exception as e:
        logger.error(f"Error al generar informe HTML: {str(e)}")
        logger.error(traceback.format_exc())

        # Crear HTML básico con mensaje de error
        return f'''
        <!DOCTYPE html>
        <html>
        <head><title>Error en informe</title></head>
        <body>
            <h1>Error al generar informe</h1>
            <p>Se produjo un error: {str(e)}</p>
            <p>Por favor, contacte con soporte técnico.</p>
        </body>
        </html>
        '''


def generar_csv_analisis(nombre, nivel, fecha, datos_analisis):
    """
    Genera un archivo CSV con los datos de análisis de una corrección.
    Versión mejorada con mejor manejo de errores y formato consistente.

    Args:
        nombre: Nombre del estudiante
        nivel: Nivel del estudiante
        fecha: Fecha de la corrección
        datos_analisis: Diccionario con datos de análisis

    Returns:
        BytesIO: Buffer con el CSV generado
    """
    try:
        # Asegurar que tenemos datos válidos
        nombre = nombre or "Estudiante"
        nivel = nivel or "No especificado"
        fecha = fecha or datetime.now().strftime("%Y-%m-%d %H:%M")

        # Verificar que datos_analisis es un diccionario
        if not isinstance(datos_analisis, dict):
            datos_analisis = {}

        # Extraer datos con manejo seguro
        errores = datos_analisis.get("errores", {}) or {}
        analisis_contextual = datos_analisis.get(
            "analisis_contextual", {}) or {}

        # Contar errores por categoría con validación
        num_gramatica = len(errores.get("Gramática", [])) if isinstance(
            errores.get("Gramática"), list) else 0
        num_lexico = len(errores.get("Léxico", [])) if isinstance(
            errores.get("Léxico"), list) else 0
        num_puntuacion = len(errores.get("Puntuación", [])) if isinstance(
            errores.get("Puntuación"), list) else 0
        num_estructura = len(errores.get("Estructura textual", [])) if isinstance(
            errores.get("Estructura textual"), list) else 0
        total_errores = num_gramatica + num_lexico + num_puntuacion + num_estructura

        # Extraer puntuaciones con validación
        coherencia = analisis_contextual.get("coherencia", {}) or {}
        cohesion = analisis_contextual.get("cohesion", {}) or {}
        registro = analisis_contextual.get("registro_linguistico", {}) or {}
        adecuacion = analisis_contextual.get("adecuacion_cultural", {}) or {}

        # Convertir a valores numéricos o 0 si no están disponibles
        def safe_numeric(val, default=0):
            if val is None:
                return default
            try:
                return float(val)
            except (ValueError, TypeError):
                return default

        coherencia_punt = safe_numeric(coherencia.get("puntuacion"))
        cohesion_punt = safe_numeric(cohesion.get("puntuacion"))
        registro_punt = safe_numeric(registro.get("puntuacion"))
        adecuacion_punt = safe_numeric(adecuacion.get("puntuacion"))

        # Extraer consejo final
        consejo_final = datos_analisis.get("consejo_final", "")
        # Limitar a 100 caracteres para el CSV
        consejo_resumen = consejo_final[:100] + "..." if consejo_final and len(
            consejo_final) > 100 else consejo_final

        # Crear CSV en memoria
        csv_buffer = StringIO()

        # Encabezados
        csv_buffer.write("Categoría,Dato\n")

        # Datos básicos
        csv_buffer.write(f"Nombre,{nombre}\n")
        csv_buffer.write(f"Nivel,{nivel}\n")
        csv_buffer.write(f"Fecha,{fecha}\n")

        # Datos de errores
        csv_buffer.write(f"Errores Gramática,{num_gramatica}\n")
        csv_buffer.write(f"Errores Léxico,{num_lexico}\n")
        csv_buffer.write(f"Errores Puntuación,{num_puntuacion}\n")
        csv_buffer.write(f"Errores Estructura,{num_estructura}\n")
        csv_buffer.write(f"Total Errores,{total_errores}\n")

        # Datos de análisis contextual
        csv_buffer.write(f"Puntuación Coherencia,{coherencia_punt}\n")
        csv_buffer.write(f"Puntuación Cohesión,{cohesion_punt}\n")
        csv_buffer.write(f"Puntuación Registro,{registro_punt}\n")
        csv_buffer.write(f"Puntuación Adecuación Cultural,{adecuacion_punt}\n")

        # Datos adicionales (sin incluir texto completo)
        csv_buffer.write(
            f"Tipo Registro,{registro.get('tipo_detectado', 'No especificado')}\n")
        csv_buffer.write(f"Consejo Final (Resumen),{consejo_resumen}\n")

        # Convertir a bytes
        csv_bytes = csv_buffer.getvalue().encode(
            'utf-8-sig')  # Usar UTF-8 con BOM para Excel

        # Crear buffer de bytes
        bytes_buffer = BytesIO(csv_bytes)
        bytes_buffer.seek(0)  # Importante: posicionar al inicio del buffer

        return bytes_buffer

    except Exception as e:
        logger.error(f"Error al generar CSV: {str(e)}")
        logger.error(traceback.format_exc())

        # Crear CSV básico con mensaje de error
        csv_buffer = StringIO()
        csv_buffer.write("Error,Mensaje\n")
        csv_buffer.write(f"Error al generar CSV,{str(e)}\n")

        return BytesIO(csv_buffer.getvalue().encode('utf-8-sig'))
    """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 7: Funciones de Procesamiento que Dependen de UI (continuación 2)
==================================================================================

Continuación de las funciones de procesamiento que dependen de la UI
"""

# --- 2. VISUALIZACIÓN DE DATOS Y ESTADÍSTICAS ---


def crear_grafico_radar(valores, categorias):
    """
    Crea un gráfico de radar para visualizar habilidades contextuales.

    Args:
        valores: Lista de valores numéricos
        categorias: Lista de nombres de categorías

    Returns:
        matplotlib.figure.Figure: Figura con el gráfico
    """
    try:
        # Verificar entradas
        if not isinstance(valores, list) or not isinstance(categorias, list):
            logger.error("Valores o categorías no son listas")
            return None

        if len(valores) != len(categorias) or len(valores) == 0:
            logger.error(
                f"Longitud incorrecta: valores={len(valores)}, categorias={len(categorias)}")
            return None

        # Convertir valores a tipo numérico
        valores_num = []
        for val in valores:
            try:
                valores_num.append(float(val))
            except (ValueError, TypeError):
                valores_num.append(0.0)

        # Crear figura
        fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))

        # Número de categorías
        N = len(categorias)

        # Ángulos para cada eje
        angulos = [n / float(N) * 2 * 3.14159 for n in range(N)]
        angulos += angulos[:1]  # Cerrar el círculo

        # Añadir los valores, repitiendo el primero
        valores_radar = valores_num + [valores_num[0]]

        # Dibujar los ejes
        plt.xticks(angulos[:-1], categorias)

        # Dibujar el polígono
        ax.plot(angulos, valores_radar)
        ax.fill(angulos, valores_radar, alpha=0.1)

        # Ajustar escala
        ax.set_yticks([2, 4, 6, 8, 10])
        ax.set_ylim(0, 10)

        plt.title("Habilidades contextuales")

        return fig

    except Exception as e:
        logger.error(f"Error al crear gráfico radar: {str(e)}")
        return None


def mostrar_progreso(df):
    """
    Crea gráficos de progreso a partir de un DataFrame con historial de correcciones.

    Args:
        df: DataFrame con historial de correcciones

    Returns:
        dict: Diccionario con gráficos generados (altair)
    """
    resultado = {
        "errores_totales": None,
        "tipos_error": None,
        "radar": None,
        "fecha_col": None
    }

    if df is None or df.empty:
        logger.warning("DataFrame vacío o nulo")
        return resultado

    try:
        # Verificar si existe la columna Fecha
        fecha_col = None
        # Buscar la columna de fecha de manera flexible
        for col in df.columns:
            if 'fecha' in col.lower().strip():
                fecha_col = col
                break

        if fecha_col is None:
            logger.error("No se encontró la columna 'Fecha' en los datos")
            return resultado

        resultado["fecha_col"] = fecha_col

        # Asegurarse de que la columna Fecha está en formato datetime
        try:
            df[fecha_col] = pd.to_datetime(df[fecha_col], errors='coerce')
            df = df.sort_values(fecha_col)
        except Exception as e:
            logger.error(f"Error al convertir fechas: {str(e)}")
            return resultado

        # Verificar que Total Errores es numérico
        if 'Total Errores' in df.columns:
            # Convertir a numérico de manera segura
            df['Total Errores'] = pd.to_numeric(
                df['Total Errores'], errors='coerce').fillna(0)

            # Crear un gráfico con Altair para total de errores
            # Convertir a formato largo para Altair
            source = pd.DataFrame({
                'Fecha': df[fecha_col],
                'Total Errores': df['Total Errores'],
                'Nivel': df.get('Nivel', 'No especificado')
            })

            chart_errores = alt.Chart(source).mark_line(point=True).encode(
                x=alt.X('Fecha:T', title='Fecha'),
                y=alt.Y('Total Errores:Q', title='Total Errores'),
                tooltip=['Fecha:T', 'Total Errores:Q', 'Nivel:N']
            ).properties(
                title='Evolución de errores totales a lo largo del tiempo'
            ).interactive()

            resultado["errores_totales"] = chart_errores

        # Gráfico de tipos de errores
        columnas_errores = [
            'Errores Gramática', 'Errores Léxico', 'Errores Puntuación',
            'Errores Estructura'
        ]

        # Encontrar las columnas que realmente existen en el DataFrame
        columnas_errores_existentes = [
            col for col in columnas_errores if col in df.columns]

        if columnas_errores_existentes:
            # Convertir columnas a numérico
            for col in columnas_errores_existentes:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)

            # Crear DataFrame en formato largo para Altair
            tipos_error_data = []
            for idx, row in df.iterrows():
                for col in columnas_errores_existentes:
                    tipos_error_data.append({
                        'Fecha': row[fecha_col],
                        'Tipo de Error': col,
                        'Cantidad': row[col]
                    })

            tipos_error_df = pd.DataFrame(tipos_error_data)

            chart_tipos = alt.Chart(tipos_error_df).mark_line(point=True).encode(
                x=alt.X('Fecha:T', title='Fecha'),
                y=alt.Y('Cantidad:Q', title='Cantidad'),
                color=alt.Color('Tipo de Error:N', title='Tipo de Error'),
                tooltip=['Fecha:T', 'Tipo de Error:N', 'Cantidad:Q']
            ).properties(
                title='Evolución por tipo de error'
            ).interactive()

            resultado["tipos_error"] = chart_tipos

        # Datos para el gráfico de radar (última entrada)
        if 'Puntuación Coherencia' in df.columns and len(df) > 0:
            ultima_entrada = df.iloc[-1]

            # Datos para el gráfico de radar
            categorias = ['Coherencia', 'Cohesión', 'Registro', 'Ad. Cultural']
            valores = [
                ultima_entrada.get('Puntuación Coherencia', 0),
                ultima_entrada.get('Puntuación Cohesión', 0),
                ultima_entrada.get('Puntuación Registro', 0),
                ultima_entrada.get('Puntuación Adecuación Cultural', 0)
            ]

            fig_radar = crear_grafico_radar(valores, categorias)
            resultado["radar"] = fig_radar

        return resultado

    except Exception as e:
        logger.error(f"Error al mostrar progreso: {str(e)}")
        return resultado
    """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 7: Funciones de Procesamiento que Dependen de UI (continuación 3)
==================================================================================

Continuación de las funciones de procesamiento que dependen de la UI
"""

# --- 3. BASE DE DATOS DE RECURSOS EDUCATIVOS ---


# Base de datos simplificada de recursos por niveles y categorías
RECURSOS_DB = {
    "A1-A2": {
        "Gramática": [
            {"título": "Presente de indicativo", "tipo": "Ficha",
                "url": "https://www.profedeele.es/gramatica/presente-indicativo/", "nivel": "A1"},
            {"título": "Los artículos en español", "tipo": "Vídeo",
                "url": "https://www.youtube.com/watch?v=example1", "nivel": "A1"},
            {"título": "Ser y estar", "tipo": "Ejercicios",
                "url": "https://aprenderespanol.org/ejercicios/ser-estar", "nivel": "A2"},
            {"título": "Pretérito indefinido", "tipo": "Explicación",
                "url": "https://www.cervantes.es/gramatica/indefinido", "nivel": "A2"}
        ],
        "Léxico": [
            {"título": "Vocabulario básico", "tipo": "Ficha",
                "url": "https://www.spanishdict.com/vocabulario-basico", "nivel": "A1"},
            {"título": "Alimentos y comidas", "tipo": "Tarjetas",
                "url": "https://quizlet.com/es/alimentos", "nivel": "A1"},
            {"título": "La ciudad", "tipo": "Podcast",
                "url": "https://spanishpod101.com/la-ciudad", "nivel": "A2"}
        ],
        "Cohesión": [
            {"título": "Conectores básicos", "tipo": "Guía",
                "url": "https://www.lingolia.com/es/conectores-basicos", "nivel": "A2"},
            {"título": "Organizar ideas", "tipo": "Ejercicios",
                "url": "https://www.todo-claro.com/organizacion", "nivel": "A2"}
        ],
        "Registro": [
            {"título": "Saludos formales e informales", "tipo": "Vídeo",
                "url": "https://www.youtube.com/watch?v=example2", "nivel": "A1"},
            {"título": "Peticiones corteses", "tipo": "Diálogos",
                "url": "https://www.lingoda.com/es/cortesia", "nivel": "A2"}
        ]
    },
    "B1-B2": {
        "Gramática": [
            {"título": "Subjuntivo presente", "tipo": "Guía",
                "url": "https://www.profedeele.es/subjuntivo-presente/", "nivel": "B1"},
            {"título": "Estilo indirecto", "tipo": "Ejercicios",
                "url": "https://www.cervantes.es/estilo-indirecto", "nivel": "B2"}
        ],
        "Léxico": [
            {"título": "Expresiones idiomáticas", "tipo": "Podcast",
                "url": "https://spanishpod101.com/expresiones", "nivel": "B1"},
            {"título": "Vocabulario académico", "tipo": "Glosario",
                "url": "https://cvc.cervantes.es/vocabulario-academico", "nivel": "B2"}
        ],
        "Cohesión": [
            {"título": "Marcadores discursivos", "tipo": "Guía",
                "url": "https://www.cervantes.es/marcadores", "nivel": "B1"},
            {"título": "Conectores argumentativos", "tipo": "Ejercicios",
                "url": "https://www.todo-claro.com/conectores", "nivel": "B2"}
        ],
        "Registro": [
            {"título": "Lenguaje formal e informal", "tipo": "Curso",
                "url": "https://www.coursera.org/spanish-registers", "nivel": "B1"},
            {"título": "Comunicación profesional", "tipo": "Ejemplos",
                "url": "https://www.cervantes.es/comunicacion-profesional", "nivel": "B2"}
        ]
    },
    "C1-C2": {
        "Gramática": [
            {"título": "Construcciones pasivas", "tipo": "Análisis",
                "url": "https://www.profedeele.es/pasivas-avanzadas/", "nivel": "C1"},
            {"título": "Subordinadas complejas", "tipo": "Guía",
                "url": "https://www.cervantes.es/subordinadas", "nivel": "C2"}
        ],
        "Léxico": [
            {"título": "Lenguaje académico", "tipo": "Corpus",
                "url": "https://www.rae.es/corpus-academico", "nivel": "C1"},
            {"título": "Variantes dialectales", "tipo": "Curso",
                "url": "https://www.coursera.org/variantes-espanol", "nivel": "C2"}
        ],
        "Cohesión": [
            {"título": "Estructura textual avanzada", "tipo": "Manual",
                "url": "https://www.uned.es/estructura-textual", "nivel": "C1"},
            {"título": "Análisis del discurso", "tipo": "Investigación",
                "url": "https://cvc.cervantes.es/analisis-discurso", "nivel": "C2"}
        ],
        "Registro": [
            {"título": "Pragmática intercultural", "tipo": "Seminario",
                "url": "https://www.cervantes.es/pragmatica", "nivel": "C1"},
            {"título": "Lenguaje literario", "tipo": "Análisis",
                "url": "https://www.rae.es/lenguaje-literario", "nivel": "C2"}
        ]
    }
}

# --- 4. GENERACIÓN DE EJERCICIOS PERSONALIZADOS ---


def generar_ejercicios_personalizado(errores_obj, analisis_contextual, nivel, idioma):
    """
    Genera ejercicios personalizados basados en los errores y análisis del estudiante.

    Args:
        errores_obj: Objeto con errores detectados
        analisis_contextual: Objeto con análisis contextual
        nivel: Nivel del estudiante
        idioma: Idioma de las instrucciones

    Returns:
        dict: Datos de ejercicios generados
    """
    client = get_openai_client()
    if client is None:
        return {"ejercicios": [{"titulo": "Servicio no disponible",
                                "tipo": "Error",
                                "instrucciones": "El servicio de generación de ejercicios no está disponible en este momento.",
                                "contenido": "Inténtelo más tarde.",
                                "solucion": "N/A"}]}

    if not circuit_breaker.can_execute("openai"):
        return {"ejercicios": [{"titulo": "Servicio temporalmente no disponible",
                                "tipo": "Error",
                                "instrucciones": "El servicio está temporalmente deshabilitado debido a errores previos.",
                                "contenido": "Inténtelo más tarde.",
                                "solucion": "N/A"}]}

    try:
        # Verificar entradas
        if not isinstance(errores_obj, dict):
            errores_obj = {}
        if not isinstance(analisis_contextual, dict):
            analisis_contextual = {}

        # Preparar datos para el prompt
        errores_gramatica = errores_obj.get("Gramática", [])
        errores_lexico = errores_obj.get("Léxico", [])
        errores_puntuacion = errores_obj.get("Puntuación", [])
        errores_estructura = errores_obj.get("Estructura textual", [])

        # Extraer puntos débiles del análisis contextual
        coherencia = analisis_contextual.get("coherencia", {})
        cohesion = analisis_contextual.get("cohesion", {})
        registro = analisis_contextual.get("registro_linguistico", {})

        # Mapear nivel para el prompt
        if "principiante" in nivel.lower():
            nivel_prompt = "A1-A2"
        elif "intermedio" in nivel.lower():
            nivel_prompt = "B1-B2"
        else:
            nivel_prompt = "C1-C2"

        # Obtener ejemplos de errores (con manejo seguro)
        ejemplos_gramatica = ", ".join([e.get('fragmento_erroneo', '')
                                       for e in errores_gramatica[:2]]) if errores_gramatica else ""
        ejemplos_lexico = ", ".join([e.get('fragmento_erroneo', '')
                                    for e in errores_lexico[:2]]) if errores_lexico else ""

        # Construir prompt para OpenAI
        prompt_ejercicios = f"""
        Basándote en los errores y análisis contextual de un estudiante de español de nivel {nivel_prompt},
        crea 3 ejercicios personalizados que le ayuden a mejorar. El estudiante tiene:
        
        - Errores gramaticales: {len(errores_gramatica)} {f"(ejemplos: {ejemplos_gramatica})" if ejemplos_gramatica else ""}
        - Errores léxicos: {len(errores_lexico)} {f"(ejemplos: {ejemplos_lexico})" if ejemplos_lexico else ""}
        - Errores de puntuación: {len(errores_puntuacion)}
        - Errores de estructura: {len(errores_estructura)}

        - Puntuación en coherencia: {coherencia.get('puntuacion', 0)}/10
        - Puntuación en cohesión: {cohesion.get('puntuacion', 0)}/10
        - Registro lingüístico: {registro.get('tipo_detectado', 'No especificado')}

        Crea ejercicios breves y específicos en formato JSON con esta estructura:
        {{
          "ejercicios": [
            {{
              "titulo": "Título del ejercicio",
              "tipo": "tipo de ejercicio (completar huecos, ordenar frases, etc.)",
              "instrucciones": "instrucciones claras y breves",
              "contenido": "el contenido del ejercicio",
              "solucion": "la solución del ejercicio"
            }}
          ]
        }}
        """

        # Idioma para las instrucciones
        if idioma != "Español":
            prompt_ejercicios += f"\nTraduce las instrucciones y el título al {idioma}, pero mantén el contenido del ejercicio en español."

        def send_request():
            return client.chat.completions.create(
                model="gpt-4-turbo",
                temperature=0.7,
                response_format={"type": "json_object"},
                messages=[{"role": "system", "content": "Eres un experto profesor de ELE especializado en crear ejercicios personalizados."},
                          {"role": "user", "content": prompt_ejercicios}]
            )

        # Usar sistema de reintentos
        response = retry_with_backoff(send_request, max_retries=2)
        raw_output = response.choices[0].message.content

        # Extraer JSON
        ejercicios_data = extract_json_safely(raw_output)

        # Verificar si se obtuvo un resultado válido
        if "error" in ejercicios_data:
            logger.warning(
                f"Error al extraer JSON de ejercicios: {ejercicios_data['error']}")
            return {"ejercicios": [{"titulo": "Ejercicio de repaso",
                                   "tipo": "Ejercicio de práctica",
                                    "instrucciones": "Revisa los elementos más problemáticos en tu texto",
                                    "contenido": "Contenido genérico de práctica",
                                    "solucion": "Consulta con tu profesor"}]}

        # Registrar éxito
        circuit_breaker.record_success("openai")
        return ejercicios_data

    except Exception as e:
        handle_exception("generar_ejercicios_personalizado", e)
        circuit_breaker.record_failure("openai")
        return {"ejercicios": [{"titulo": "Error en la generación",
                                "tipo": "Error controlado",
                                "instrucciones": "No se pudieron generar ejercicios personalizados",
                                "contenido": f"Error: {str(e)}",
                                "solucion": "Intenta de nuevo más tarde"}]}


def obtener_recursos_recomendados(errores_obj, analisis_contextual, nivel):
    """
    Obtiene recursos recomendados basados en los errores y análisis del estudiante.

    Args:
        errores_obj: Objeto con errores detectados
        analisis_contextual: Objeto con análisis contextual
        nivel: Nivel del estudiante

    Returns:
        list: Lista de recursos recomendados
    """
    recursos_recomendados = []

    try:
        # Verificar entradas
        if not isinstance(errores_obj, dict):
            errores_obj = {}
        if not isinstance(analisis_contextual, dict):
            analisis_contextual = {}

        # Determinar el nivel para buscar recursos
        if "principiante" in nivel.lower():
            nivel_db = "A1-A2"
        elif "intermedio" in nivel.lower():
            nivel_db = "B1-B2"
        else:
            nivel_db = "C1-C2"

        # Verificar errores gramaticales
        if len(errores_obj.get("Gramática", [])) > 0:
            recursos_gramatica = RECURSOS_DB.get(
                nivel_db, {}).get("Gramática", [])
            if recursos_gramatica:
                recursos_recomendados.extend(recursos_gramatica[:2])

        # Verificar errores léxicos
        if len(errores_obj.get("Léxico", [])) > 0:
            recursos_lexico = RECURSOS_DB.get(nivel_db, {}).get("Léxico", [])
            if recursos_lexico:
                recursos_recomendados.extend(recursos_lexico[:2])

        # Verificar problemas de cohesión
        if analisis_contextual.get("cohesion", {}).get("puntuacion", 10) < 7:
            recursos_cohesion = RECURSOS_DB.get(
                nivel_db, {}).get("Cohesión", [])
            if recursos_cohesion:
                recursos_recomendados.extend(recursos_cohesion[:1])

        # Verificar problemas de registro
        if analisis_contextual.get("registro_linguistico", {}).get("puntuacion", 10) < 7:
            recursos_registro = RECURSOS_DB.get(
                nivel_db, {}).get("Registro", [])
            if recursos_registro:
                recursos_recomendados.extend(recursos_registro[:1])

    except Exception as e:
        logger.error(f"Error al obtener recursos: {str(e)}")

    return recursos_recomendados


"""
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 7: Funciones de Procesamiento que Dependen de UI (continuación 4)
==================================================================================

Continuación de las funciones de procesamiento que dependen de la UI
"""

# --- 5. GENERACIÓN DE PLAN DE ESTUDIO PERSONALIZADO ---


def generar_plan_estudio_personalizado(nombre, nivel, datos_historial):
    """
    Genera un plan de estudio personalizado basado en el historial del estudiante.

    Args:
        nombre: Nombre del estudiante
        nivel: Nivel del estudiante
        datos_historial: DataFrame con historial de correcciones

    Returns:
        dict: Plan de estudio generado
    """
    client = get_openai_client()
    if client is None:
        return {"error": "Servicio no disponible", "plan": None}

    if not circuit_breaker.can_execute("openai"):
        return {"error": "Servicio temporalmente no disponible", "plan": None}

    if datos_historial is None or datos_historial.empty:
        return {"error": "No hay suficientes datos para generar un plan personalizado", "plan": None}

    try:
        # Extraer estadísticas básicas
        if 'Errores Gramática' in datos_historial.columns and 'Errores Léxico' in datos_historial.columns:
            # Calcular promedios
            promedio_gramatica = datos_historial['Errores Gramática'].mean()
            promedio_lexico = datos_historial['Errores Léxico'].mean()

            # Verificar columnas de análisis contextual
            coherencia_promedio = datos_historial['Puntuación Coherencia'].mean(
            ) if 'Puntuación Coherencia' in datos_historial.columns else 5
            cohesion_promedio = datos_historial['Puntuación Cohesión'].mean(
            ) if 'Puntuación Cohesión' in datos_historial.columns else 5

            # Extraer nivel del último registro
            if 'Nivel' in datos_historial.columns:
                nivel_actual = datos_historial.iloc[-1]['Nivel']
            else:
                nivel_actual = nivel

            # Verificar consejos finales para extraer temas recurrentes
            temas_recurrentes = []
            if 'Consejo Final' in datos_historial.columns:
                # Aquí podríamos implementar un análisis más sofisticado de los consejos
                temas_recurrentes = ["conjugación verbal",
                                     "uso de preposiciones", "concordancia"]

            # Construir contexto para la IA
            errores_frecuentes = (
                f"Promedio de errores gramaticales: {promedio_gramatica:.1f}, "
                f"Promedio de errores léxicos: {promedio_lexico:.1f}. "
                f"Puntuación en coherencia: {coherencia_promedio:.1f}/10, "
                f"Puntuación en cohesión: {cohesion_promedio:.1f}/10. "
                f"Temas recurrentes: {', '.join(temas_recurrentes)}."
            )

            # Prompt para la IA
            prompt_plan = f"""
            Crea un plan de estudio personalizado para un estudiante de español llamado {nombre} de nivel {nivel_actual}
            con los siguientes errores frecuentes: {errores_frecuentes}

            Organiza el plan por semanas (4 semanas) con objetivos claros, actividades concretas y recursos recomendados.
            Para cada semana, incluye:

            1. Objetivos específicos
            2. Temas gramaticales a trabajar
            3. Vocabulario a practicar
            4. 1-2 actividades concretas
            5. Recursos o materiales recomendados

            Adapta todo el contenido al nivel del estudiante y sus necesidades específicas.
            """

            def send_request():
                return client.chat.completions.create(
                    model="gpt-4-turbo",
                    temperature=0.7,
                    messages=[
                        {"role": "system", "content": "Eres un experto en diseño curricular ELE que crea planes de estudio personalizados."},
                        {"role": "user", "content": prompt_plan}
                    ]
                )

            # Usar sistema de reintentos
            response = retry_with_backoff(send_request, max_retries=2)
            plan_estudio = response.choices[0].message.content

            # Registrar éxito
            circuit_breaker.record_success("openai")

            # Dividir el plan por semanas
            semanas = plan_estudio.split("Semana")

            # Procesar el resultado
            plan_procesado = {
                "completo": plan_estudio,
                "semanas": []
            }

            # Ignorar el elemento vacío al inicio
            for i, semana in enumerate(semanas[1:], 1):
                titulo_semana = extraer_titulo(semana)
                contenido_semana = semana.strip()
                plan_procesado["semanas"].append({
                    "numero": i,
                    "titulo": titulo_semana,
                    "contenido": contenido_semana
                })

            return {"error": None, "plan": plan_procesado}

        else:
            return {"error": "No se encontraron columnas necesarias en los datos", "plan": None}

    except Exception as e:
        handle_exception("generar_plan_estudio_personalizado", e)
        circuit_breaker.record_failure("openai")
        return {"error": f"Error al generar plan de estudio: {str(e)}", "plan": None}

    """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 8: Componentes de UI Avanzados y Pestañas
==================================================================================

Este artefacto contiene:
1. Componentes UI avanzados que dependen de funciones previas
2. Visualizadores y exportadores de resultados
3. Implementación de pestañas principales
4. Implementación de subpestañas
"""

# --- 1. COMPONENTES UI AVANZADOS ---


def ui_show_correction_results(result, show_export=True):
    """
    Muestra los resultados de una corrección de texto.

    Args:
        result: Resultados de la corrección
        show_export: Mostrar opciones de exportación
    """
    if "error" in result:
        st.error(f"Error en la corrección: {result['error']}")
        return

    # Extraer campos del JSON
    saludo = result.get("saludo", "")
    tipo_texto_detectado = result.get("tipo_texto", "")
    errores_obj = result.get("errores", {})
    texto_corregido = result.get("texto_corregido", "")
    analisis_contextual = result.get("analisis_contextual", {})
    consejo_final = result.get("consejo_final", "")
    fin = result.get("fin", "")

    # Extraer puntuaciones del análisis contextual
    coherencia = analisis_contextual.get("coherencia", {})
    cohesion = analisis_contextual.get("cohesion", {})
    registro = analisis_contextual.get("registro_linguistico", {})
    adecuacion = analisis_contextual.get("adecuacion_cultural", {})

    puntuacion_coherencia = coherencia.get("puntuacion", 0)
    puntuacion_cohesion = cohesion.get("puntuacion", 0)
    puntuacion_registro = registro.get("puntuacion", 0)
    puntuacion_adecuacion = adecuacion.get("puntuacion", 0)

    # --- CONTEO DE ERRORES ---
    num_gramatica = len(errores_obj.get("Gramática", []))
    num_lexico = len(errores_obj.get("Léxico", []))
    num_puntuacion = len(errores_obj.get("Puntuación", []))
    num_estructura = len(errores_obj.get("Estructura textual", []))
    total_errores = num_gramatica + num_lexico + num_puntuacion + num_estructura

    # --- MOSTRAR RESULTADOS EN LA INTERFAZ ---
    # Mostrar el saludo directamente
    st.write(saludo)

    # Mostrar tipo de texto detectado con contextualización
    st.info(
        f"He identificado tu escrito como un texto de tipo **{tipo_texto_detectado.lower()}**.")

    # Errores detectados
    st.subheader("Errores detectados")
    if not any(errores_obj.get(cat, []) for cat in ["Gramática", "Léxico", "Puntuación", "Estructura textual"]):
        st.success("¡Felicidades! No se han detectado errores significativos.")
    else:
        for categoria in ["Gramática", "Léxico", "Puntuación", "Estructura textual"]:
            lista_errores = errores_obj.get(categoria, [])
            if lista_errores:
                with st.expander(f"**{categoria}** ({len(lista_errores)} errores)"):
                    for i, err in enumerate(lista_errores, 1):
                        st.markdown(f"**Error {i}:**")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.error(f"❌ {err.get('fragmento_erroneo', '')}")
                        with col2:
                            st.success(f"✅ {err.get('correccion', '')}")
                        st.info(f"💡 {err.get('explicacion', '')}")
                        if i < len(lista_errores):
                            st.divider()

    # Texto corregido
    st.subheader("Texto corregido completo")
    st.write(texto_corregido)

    # --- ANÁLISIS CONTEXTUAL ---
    st.header("Análisis contextual avanzado")

    # Crear columnas para las puntuaciones generales
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Coherencia", f"{puntuacion_coherencia}/10")
    with col2:
        st.metric("Cohesión", f"{puntuacion_cohesion}/10")
    with col3:
        st.metric("Registro", f"{puntuacion_registro}/10")
    with col4:
        st.metric("Adecuación cultural", f"{puntuacion_adecuacion}/10")

    # Gráfico sencillo para visualizar las puntuaciones
    puntuaciones = [puntuacion_coherencia, puntuacion_cohesion,
                    puntuacion_registro, puntuacion_adecuacion]
    categorias = ["Coherencia", "Cohesión", "Registro", "Ad. Cultural"]

    # Calcular el promedio de las puntuaciones
    promedio_contextual = sum(puntuaciones) / \
        len(puntuaciones) if puntuaciones else 0

    # Mostrar un progreso general
    st.markdown(f"##### Evaluación global: {promedio_contextual:.1f}/10")
    st.progress(promedio_contextual / 10)

    # Detalles de coherencia
    with st.expander("Coherencia textual", expanded=True):
        st.markdown(f"**Comentario**: {coherencia.get('comentario', '')}")
        sugerencias = coherencia.get("sugerencias", [])
        if sugerencias:
            st.markdown("**Sugerencias para mejorar:**")
            for sug in sugerencias:
                st.markdown(f"- {sug}")

    # Detalles de cohesión
    with st.expander("Cohesión textual", expanded=True):
        st.markdown(f"**Comentario**: {cohesion.get('comentario', '')}")
        sugerencias = cohesion.get("sugerencias", [])
        if sugerencias:
            st.markdown("**Sugerencias para mejorar:**")
            for sug in sugerencias:
                st.markdown(f"- {sug}")

    # Detalles de registro lingüístico
    with st.expander("Registro lingüístico", expanded=True):
        st.markdown(
            f"**Tipo de registro detectado**: {registro.get('tipo_detectado', '')}")
        st.markdown(
            f"**Adecuación al contexto**: {registro.get('adecuacion', '')}")
        sugerencias = registro.get("sugerencias", [])
        if sugerencias:
            st.markdown("**Sugerencias para mejorar:**")
            for sug in sugerencias:
                st.markdown(f"- {sug}")

    # Detalles de adecuación cultural
    with st.expander("Adecuación cultural y pragmática", expanded=True):
        st.markdown(f"**Comentario**: {adecuacion.get('comentario', '')}")
        elementos = adecuacion.get("elementos_destacables", [])
        if elementos:
            st.markdown("**Elementos culturales destacables:**")
            for elem in elementos:
                st.markdown(f"- {elem}")
        sugerencias = adecuacion.get("sugerencias", [])
        if sugerencias:
            st.markdown("**Sugerencias para mejorar:**")
            for sug in sugerencias:
                st.markdown(f"- {sug}")

    # Consejo final
    st.subheader("Consejo final")
    st.info(consejo_final)
    st.write(fin)

    # --- GENERAR AUDIO CON ELEVENLABS (Consejo final en español) ---
    if consejo_final:
        st.markdown("**🔊 Consejo leído en voz alta:**")
        with st.spinner("Generando audio con ElevenLabs..."):
            audio_bytes = generar_audio_consejo(consejo_final)
            if audio_bytes:
                st.audio(audio_bytes, format="audio/mpeg")
            else:
                st.warning("⚠️ No se pudo generar el audio del consejo.")

    # --- MOSTRAR RECOMENDACIONES PERSONALIZADAS ---
    ui_show_recommendations(errores_obj, analisis_contextual, get_session_var(
        "nivel_estudiante", "intermedio"), "Spanish")

    # --- OPCIONES DE EXPORTACIÓN ---
    if show_export:
        ui_export_options(result)
        """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 8: Componentes de UI Avanzados y Pestañas (continuación 1)
==================================================================================

Continuación de los componentes de UI avanzados y pestañas
"""


def ui_show_recommendations(errores_obj, analisis_contextual, nivel, idioma):
    """
    Muestra recomendaciones personalizadas basadas en el análisis.

    Args:
        errores_obj: Objeto con errores detectados
        analisis_contextual: Objeto con análisis contextual
        nivel: Nivel del estudiante
        idioma: Idioma de las instrucciones
    """
    st.header("📚 Recomendaciones personalizadas")

    # Pestañas para diferentes tipos de recomendaciones
    tab1, tab2 = st.tabs(
        ["📖 Recursos recomendados", "✏️ Ejercicios personalizados"])

    with tab1:
        recursos = obtener_recursos_recomendados(
            errores_obj, analisis_contextual, nivel)

        if recursos:
            st.write("Basado en tu análisis, te recomendamos estos recursos:")

            for i, recurso in enumerate(recursos):
                col1, col2, col3 = st.columns([2, 1, 1])
                with col1:
                    st.markdown(f"**{recurso['título']}**")
                with col2:
                    st.write(f"Tipo: {recurso['tipo']}")
                with col3:
                    st.write(f"Nivel: {recurso['nivel']}")
                st.markdown(f"[Ver recurso]({recurso['url']})")
                if i < len(recursos) - 1:
                    st.divider()
        else:
            st.info("No hay recursos específicos para recomendar en este momento.")

    with tab2:
        st.write("Ejercicios personalizados según tus necesidades:")

        with st.spinner("Generando ejercicios personalizados..."):
            ejercicios_data = generar_ejercicios_personalizado(
                errores_obj, analisis_contextual, nivel, idioma)

            ejercicios = ejercicios_data.get("ejercicios", [])

            for i, ejercicio in enumerate(ejercicios):
                # Usar st.expander para el ejercicio principal
                with st.expander(f"{ejercicio.get('titulo', f'Ejercicio {i+1}')}"):
                    # Crear pestañas para ejercicio y solución
                    ejercicio_tab, solucion_tab = st.tabs(
                        ["Ejercicio", "Solución"])

                    with ejercicio_tab:
                        st.markdown(
                            f"**{ejercicio.get('tipo', 'Actividad')}**")
                        st.markdown(
                            f"*Instrucciones:* {ejercicio.get('instrucciones', '')}")
                        st.markdown("---")
                        st.markdown(ejercicio.get('contenido', ''))

                    with solucion_tab:
                        st.markdown(f"#### Solución del ejercicio:")
                        st.markdown(ejercicio.get('solucion', ''))


def ui_export_options(data):
    """
    Muestra opciones para exportar los resultados de la corrección.
    Versión mejorada con mejor manejo de errores y comprobaciones.

    Args:
        data: Resultados de la corrección
    """
    st.header("📊 Exportar informe")

    # Verificar que existen campos necesarios
    if not isinstance(data, dict) or "texto_corregido" not in data:
        st.warning("⚠️ No hay datos suficientes para exportar.")
        return

    # Extraer datos para la exportación con manejo seguro
    nombre = get_session_var("usuario_actual", "Usuario")
    nivel = get_session_var("nivel_estudiante", "intermedio")
    fecha = datetime.now().strftime("%Y-%m-%d %H:%M")
    texto_original = get_session_var("ultimo_texto", "")
    texto_corregido = data.get("texto_corregido", "")
    errores_obj = data.get("errores", {})
    analisis_contextual = data.get("analisis_contextual", {})
    consejo_final = data.get("consejo_final", "")

    # Opciones de exportación en pestañas
    export_tab1, export_tab2, export_tab3 = st.tabs(
        ["📝 Documento Word", "🌐 Documento HTML", "📊 Excel/CSV"]
    )

    with export_tab1:
        st.write("Exporta este informe como documento Word (DOCX)")

        if st.button("Generar documento Word", key="gen_docx"):
            with st.spinner("Generando documento Word..."):
                try:
                    # Generar el documento con manejo explícito de errores
                    docx_buffer = generar_informe_docx(
                        nombre, nivel, fecha, texto_original, texto_corregido,
                        errores_obj, analisis_contextual, consejo_final
                    )

                    # Si el buffer se generó correctamente, mostrar el botón de descarga
                    if docx_buffer is not None:
                        # Verificar tamaño del buffer
                        if len(docx_buffer.getvalue()) > 0:
                            nombre_archivo = f"informe_{nombre.replace(' ', '_')}_{fecha.replace(':', '_').replace(' ', '_')}.docx"
                            st.download_button(
                                label="📥 Descargar documento Word",
                                data=docx_buffer,
                                file_name=nombre_archivo,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="docx_download_corregir",
                                help="Haz clic para descargar el informe en formato Word"
                            )
                            st.success("✅ Documento generado correctamente")
                        else:
                            st.error(
                                "El documento generado está vacío. Inténtalo de nuevo.")
                    else:
                        st.error(
                            "No se pudo generar el documento Word. Verifica que todas las dependencias estén instaladas.")
                except Exception as e:
                    st.error(f"Error al generar el documento Word: {str(e)}")
                    logger.error(f"Error en generación de DOCX: {str(e)}")
                    logger.error(traceback.format_exc())

    with export_tab2:
        st.write("Exporta este informe como página web (HTML)")

        if st.button("Generar documento HTML", key="gen_html"):
            with st.spinner("Generando HTML..."):
                try:
                    # Generar el HTML
                    html_content = generar_informe_html(
                        nombre, nivel, fecha, texto_original, texto_corregido,
                        analisis_contextual, consejo_final
                    )

                    # Convertir a bytes para descargar
                    html_bytes = html_content.encode()

                    # Verificar que tenemos contenido
                    if len(html_bytes) > 0:
                        # Botón de descarga
                        nombre_archivo = f"informe_{nombre.replace(' ', '_')}_{fecha.replace(':', '_').replace(' ', '_')}.html"
                        st.download_button(
                            label="📥 Descargar página HTML",
                            data=html_bytes,
                            file_name=nombre_archivo,
                            mime="text/html",
                            key="html_download_corregir",
                            help="Haz clic para descargar el informe como página web HTML"
                        )
                        st.success("✅ HTML generado correctamente")

                        # Opción para previsualizar
                        with st.expander("Previsualizar HTML"):
                            # Sanitizar para evitar problemas con comillas
                            sanitized_html = html_content.replace(
                                '"', '&quot;')
                            st.markdown(
                                f'<iframe srcdoc="{sanitized_html}" width="100%" height="600" style="border: 1px solid #ddd; border-radius: 5px;"></iframe>',
                                unsafe_allow_html=True
                            )
                    else:
                        st.error(
                            "El HTML generado está vacío. Inténtalo de nuevo.")
                except Exception as e:
                    st.error(f"Error al generar el HTML: {str(e)}")
                    logger.error(f"Error en generación de HTML: {str(e)}")
                    logger.error(traceback.format_exc())

    with export_tab3:
        st.write("Exporta los datos del análisis en formato CSV")

        if st.button("Generar CSV", key="gen_csv"):
            with st.spinner("Generando CSV..."):
                try:
                    # Generar el CSV
                    csv_buffer = generar_csv_analisis(
                        nombre, nivel, fecha, data
                    )

                    # Verificar que tenemos contenido
                    if csv_buffer and len(csv_buffer.getvalue()) > 0:
                        # Botón de descarga
                        nombre_archivo = f"datos_{nombre.replace(' ', '_')}_{fecha.replace(':', '_').replace(' ', '_')}.csv"
                        st.download_button(
                            label="📥 Descargar CSV",
                            data=csv_buffer,
                            file_name=nombre_archivo,
                            mime="text/csv",
                            key="csv_download_corregir",
                            help="Haz clic para descargar los datos de análisis en formato CSV"
                        )
                        st.success("✅ CSV generado correctamente")
                    else:
                        st.error(
                            "El CSV generado está vacío. Inténtalo de nuevo.")
                except Exception as e:
                    st.error(f"Error al generar el CSV: {str(e)}")
                    logger.error(f"Error en generación de CSV: {str(e)}")
                    logger.error(traceback.format_exc())
                    """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 8: Componentes de UI Avanzados y Pestañas (continuación 2)
==================================================================================

Continuación de los componentes de UI avanzados y pestañas
"""

# --- 2. IMPLEMENTACIÓN DE PESTAÑAS PRINCIPALES ---


def tab_corregir():
    """Implementación de la pestaña de corrección de texto."""
    # Limpiar variables de simulacro si venimos de otra pestaña
    if "inicio_simulacro" in st.session_state:
        set_session_var("inicio_simulacro", None)
    if "duracion_simulacro" in st.session_state:
        set_session_var("duracion_simulacro", None)
    if "tarea_simulacro" in st.session_state:
        set_session_var("tarea_simulacro", None)

    st.header("📝 Corrección de texto")

    with st.expander("ℹ️ Información sobre el análisis contextual", expanded=False):
        st.markdown("""
        Esta versión mejorada del Textocorrector incluye:
        - **Análisis de coherencia**: Evalúa si las ideas están conectadas de manera lógica y si el texto tiene sentido en su conjunto.
        - **Análisis de cohesión**: Revisa los mecanismos lingüísticos que conectan las diferentes partes del texto.
        - **Evaluación del registro lingüístico**: Determina si el lenguaje usado es apropiado para el contexto y propósito del texto.
        - **Análisis de adecuación cultural**: Identifica si hay expresiones o referencias culturalmente apropiadas o inapropiadas.

        Las correcciones se adaptan automáticamente al nivel del estudiante.
        """)

    # Obtener datos del usuario - usando key única para este formulario
    user_data = ui_user_info_form(form_key="form_user_info_corregir")

    # El resto del código permanece igual...
    if not user_data:
        if "usuario_actual" not in st.session_state or not st.session_state.usuario_actual:
            st.info("👆 Por favor, introduce tu nombre y nivel para comenzar.")
            return

    # GENERADOR DE CONSIGNAS
    with st.expander("¿No sabes qué escribir? Yo te ayudo...", expanded=False):
        tipo_consigna = st.selectbox(
            "Tipo de texto a escribir:",
            [
                "Cualquiera (aleatorio)",
                "Narración",
                "Correo/Carta formal",
                "Opinión/Argumentación",
                "Descripción",
                "Diálogo"
            ],
            key="tipo_consigna_corregir"
        )

        if st.button("Generar consigna de escritura", key="generar_consigna"):
            with st.spinner("Generando consigna adaptada a tu nivel..."):
                # Determinar el nivel para la IA
                nivel_actual = get_session_var(
                    "nivel_estudiante", "intermedio")

                # Generar la consigna
                consigna_generada = generar_consigna_escritura(
                    nivel_actual, tipo_consigna)

                # Guardar en session_state
                set_session_var("consigna_actual", consigna_generada)

            # Mostrar la consigna generada
            if "consigna_actual" in st.session_state and st.session_state.consigna_actual:
                st.success("✨ Consigna generada:")
                st.info(st.session_state.consigna_actual)

                # Opción para usar esta consigna
                if st.button("Usar esta consigna como contexto", key="usar_consigna"):
                    set_session_var("info_adicional_corregir",
                                    f"Consigna: {st.session_state.consigna_actual}")
                    set_session_var("usar_consigna_como_texto", True)
                    st.rerun()  # Recargar para actualizar el formulario

    # FORMULARIO DE CORRECCIÓN
    with st.form(key="formulario_corregir"):
        # Opciones de corrección
        options = ui_idioma_correcciones_tipo()

        # Texto inicial con contenido de la consigna si está disponible
        texto_inicial = ""
        if get_session_var("usar_consigna_como_texto", False) and "consigna_actual" in st.session_state:
            texto_inicial = f"[Instrucción: {st.session_state.consigna_actual}]\n\n"
            # Reset para no añadirlo cada vez
            set_session_var("usar_consigna_como_texto", False)

        # Área de texto para la corrección
        texto = st.text_area(
            "Escribe tu texto aquí:",
            value=texto_inicial,
            height=250,
            key="texto_correccion_corregir"
        )

        info_adicional = st.text_area(
            "Información adicional o contexto (opcional):",
            value=get_session_var("info_adicional_corregir", ""),
            height=100,
            key="info_adicional_widget_key"  # Clave diferente para el widget
        )

        # Guardar sin conflicto (dentro del formulario, pero fuera de la definición del widget)
        set_session_var("info_adicional_corregir", info_adicional)

        # Botón de envío
        enviar = st.form_submit_button("Corregir", use_container_width=True)

        # PROCESAMIENTO DEL FORMULARIO (dentro del form, será ejecutado solo cuando se envíe)
        if enviar:
            if not texto.strip():
                st.warning("Por favor, escribe un texto para corregir.")
            else:
                # Guardar el texto para posible uso futuro
                set_session_var("ultimo_texto", texto)

                with st.spinner("Analizando texto y generando corrección contextual..."):
                    # Obtener los datos del usuario
                    nombre = get_session_var("usuario_actual", "")
                    nivel = user_data.get("nivel") if user_data else get_session_var(
                        "nivel_estudiante", "intermedio")

                    # Obtener las opciones seleccionadas
                    idioma = options.get("idioma", "Español")
                    tipo_texto = options.get(
                        "tipo_texto", "General/No especificado")
                    contexto_cultural = options.get(
                        "contexto_cultural", "General/Internacional")

                    # Llamar a la función de corrección
                    resultado = corregir_texto(
                        texto, nombre, nivel, idioma, tipo_texto, contexto_cultural, info_adicional
                    )

                    # Guardar el resultado para futuras referencias
                    set_session_var("correction_result", resultado)
                    set_session_var("last_correction_time",
                                    datetime.now().isoformat())

    # MOSTRAR RESULTADOS (fuera del form, para evitar rerun issues)
    # Comprobar si hay resultados después de enviar el formulario
    if enviar and texto.strip():
        if "error" not in resultado:
            ui_show_correction_results(resultado)
        else:
            st.error(f"Error en la corrección: {resultado['error']}")

            # Mostrar sugerencia de reintentar
            st.info(
                "Prueba a hacer la corrección con un texto más corto o inténtalo más tarde.")
            """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 8: Componentes de UI Avanzados y Pestañas (continuación 3)
==================================================================================

Continuación de los componentes de UI avanzados y pestañas
"""

# --- 3. FUNCIÓN DE VISUALIZACIÓN DE TEXTO TRANSCRITO ---


def visualizar_texto_manuscrito():
    """
    Función corregida para visualizar y corregir texto transcrito de imágenes.
    Soluciona problemas de flujo entre transcripción y corrección.
    """
    st.subheader("Corrección de texto manuscrito transcrito")

    # Verificar si hay texto transcrito para corregir
    texto_transcrito = get_session_var("ultimo_texto_transcrito", "")
    if not texto_transcrito:
        st.info("No hay texto transcrito para corregir.")
        # Botón para volver a la herramienta de transcripción
        if st.button("Volver a transcripción", key="volver_transcripcion"):
            set_session_var("mostrar_correccion_transcripcion", False)
            st.rerun()
        return

    # Mostrar texto transcrito
    texto_transcrito_editable = st.text_area(
        "Texto transcrito (puedes editarlo si hay errores):",
        value=texto_transcrito,
        height=200,
        key="texto_transcrito_editable"
    )

    # Opciones de corrección
    options = ui_idioma_correcciones_tipo()

    # Botón para enviar a corrección
    if st.button("Enviar a corrección", key="corregir_texto_transcrito_btn"):
        if not texto_transcrito_editable.strip():
            st.warning(
                "El texto está vacío. Por favor, asegúrate de que hay contenido para corregir.")
            return

        # Guardar para futura referencia
        set_session_var("ultimo_texto", texto_transcrito_editable)

        with st.spinner("Analizando texto transcrito..."):
            try:
                # Obtener datos necesarios
                nombre = get_session_var("usuario_actual", "Usuario")
                nivel = get_session_var("nivel_estudiante", "intermedio")

                # Llamar a la función de corrección
                resultado = corregir_texto(
                    texto_transcrito_editable, nombre, nivel, options["idioma"],
                    options["tipo_texto"], options["contexto_cultural"],
                    "Texto transcrito de imagen manuscrita"
                )

                # Guardar resultado
                set_session_var("correction_result", resultado)
                set_session_var("last_correction_time",
                                datetime.now().isoformat())

                # Mostrar resultados
                if "error" not in resultado:
                    ui_show_correction_results(resultado)

                    # IMPORTANTE: Botón para volver después de la corrección
                    if st.button("Volver a transcripción", key="volver_despues_correccion"):
                        set_session_var(
                            "mostrar_correccion_transcripcion", False)
                        st.rerun()
                else:
                    st.error(f"Error en la corrección: {resultado['error']}")
            except Exception as e:
                st.error(f"Error durante la corrección: {str(e)}")
                logger.error(f"Error en visualizar_texto_manuscrito: {str(e)}")
                logger.error(traceback.format_exc())

    # IMPORTANTE: Botón para cancelar y volver
    if st.button("Cancelar y volver", key="cancelar_correccion_transcripcion"):
        set_session_var("mostrar_correccion_transcripcion", False)
        st.rerun()


# --- 4. IMPLEMENTACIÓN DE SUBPESTAÑAS DE EXÁMENES ---
def modelo_examen_tab(tipo_examen, nivel_examen):
    """
    Implementación de la pestaña de modelo de examen.

    Args:
        tipo_examen: Tipo de examen seleccionado
        nivel_examen: Nivel de examen seleccionado
    """
    st.subheader("Modelo de prueba escrita")
    st.markdown("""
    Aquí encontrarás un modelo de tarea de expresión escrita similar a la que encontrarás en el examen.
    Practica sin límite de tiempo y recibe correcciones detalladas.
    """)

    # Inicialización de variables de sesión para el modelo de examen
    if "tarea_modelo_generada" not in st.session_state:
        set_session_var("tarea_modelo_generada", None)
    if "respuesta_modelo_examen" not in st.session_state:
        set_session_var("respuesta_modelo_examen", "")

    # Botón para generar tarea
    if st.button("Generar tarea de examen", key="generar_tarea_examen"):
        # Generar tarea específica para el examen y nivel seleccionados
        with st.spinner("Generando tarea oficial..."):
            tarea_generada = generar_tarea_examen(tipo_examen, nivel_examen)
            set_session_var("tarea_modelo_generada", tarea_generada)
            st.success("✅ Tarea generada correctamente")
            st.rerun()  # Refrescar para mostrar la tarea generada

    # Mostrar la tarea y área de respuesta si hay una tarea generada
    tarea_modelo = get_session_var("tarea_modelo_generada", None)
    if tarea_modelo:
        with st.expander("Ver instrucciones de la tarea", expanded=True):
            st.markdown(tarea_modelo)

        # Área para que el estudiante escriba su respuesta
        respuesta_estudiante = st.text_area(
            "Escribe tu respuesta a la tarea aquí:",
            value=get_session_var("respuesta_modelo_examen", ""),
            height=250,
            key="respuesta_modelo_examen_area"
        )

        # Guardar respuesta en session_state
        set_session_var("respuesta_modelo_examen", respuesta_estudiante)

        # Botones para opciones
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Corregir respuesta", key="finalizar_modelo_examen"):
                if respuesta_estudiante.strip():
                    # Mostrar spinner durante la corrección
                    with st.spinner("Analizando respuesta..."):
                        # Corrección integrada (sin redirección)
                        resultado = corregir_examen(
                            respuesta_estudiante,
                            tipo_examen,
                            nivel_examen
                        )

                    # Mostrar resultados directamente aquí
                    if "error" not in resultado:
                        ui_show_correction_results(resultado)
                    else:
                        st.error(
                            f"Error en la corrección: {resultado['error']}")
                else:
                    st.warning(
                        "Por favor, escribe una respuesta antes de enviar a corrección.")

        with col2:
            if st.button("Generar nueva tarea", key="nueva_tarea_modelo"):
                # Reiniciar variables
                set_session_var("tarea_modelo_generada", None)
                set_session_var("respuesta_modelo_examen", "")
                st.rerun()
                """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 8: Componentes de UI Avanzados y Pestañas (continuación 4)
==================================================================================

Continuación de los componentes de UI avanzados y pestañas
"""


def simulacro_cronometrado_tab(tipo_examen, nivel_examen):
    """
    Implementación de la pestaña de simulacro cronometrado.

    Args:
        tipo_examen: Tipo de examen seleccionado
        nivel_examen: Nivel de examen seleccionado
    """
    st.subheader("Simulacro cronometrado")
    st.markdown("""
    Pon a prueba tus habilidades bajo las condiciones reales del examen.
    Esta prueba está cronometrada según los tiempos oficiales.
    """)

    # Placeholder para el temporizador
    tiempo_restante_placeholder = st.empty()

    # Verificar si el simulacro está en progreso
    if "inicio_simulacro" not in st.session_state or get_session_var("inicio_simulacro") is None:
        if st.button("Iniciar simulacro", key="iniciar_simulacro"):
            try:
                # Configurar el temporizador con verificación
                duracion = obtener_duracion_examen(tipo_examen, nivel_examen)
                if duracion is None or duracion <= 0:
                    st.warning(
                        "No se pudo determinar la duración del examen. Usando 45 minutos por defecto.")
                    duracion = 45 * 60  # 45 minutos por defecto

                set_session_var("inicio_simulacro", time.time())
                set_session_var("duracion_simulacro", duracion)

                # Inicializar variable para la respuesta
                if "simulacro_respuesta_texto" not in st.session_state:
                    set_session_var("simulacro_respuesta_texto", "")

                # Generar tarea para el simulacro
                tarea_simulacro = generar_tarea_examen(
                    tipo_examen, nivel_examen)
                if tarea_simulacro:
                    set_session_var("tarea_simulacro", tarea_simulacro)
                else:
                    set_session_var(
                        "tarea_simulacro", "No se pudo generar la tarea. Por favor, inténtalo de nuevo.")

                st.rerun()  # Refrescar para mostrar el simulacro
            except Exception as e:
                handle_exception("iniciar_simulacro", e)
                st.error(f"Error al iniciar el simulacro: {str(e)}")
    else:
        try:
            # Simulacro en progreso
            inicio_simulacro = get_session_var("inicio_simulacro")
            duracion_simulacro = get_session_var("duracion_simulacro")

            # Verificar que los valores sean válidos
            if inicio_simulacro is None or duracion_simulacro is None:
                st.warning("Datos del simulacro incorrectos. Reiniciando...")
                set_session_var("inicio_simulacro", None)
                set_session_var("duracion_simulacro", None)
                st.rerun()

            # Obtener estado del temporizador
            timer_state = ui_countdown_timer(
                duracion_simulacro, inicio_simulacro)
        except Exception as e:
            # Manejo genérico de errores
            handle_exception("simulacro_timer", e)
            st.error(f"Error en el temporizador del simulacro: {str(e)}")
            timer_state = {"color": "error",
                           "tiempo_formateado": "00:00", "terminado": True}

        # Mostrar temporizador según el estado
        if timer_state["color"] == "normal":
            tiempo_restante_placeholder.info(
                f"⏱️ Tiempo restante: {timer_state['tiempo_formateado']}")
        elif timer_state["color"] == "warning":
            tiempo_restante_placeholder.warning(
                f"⏱️ Tiempo restante: {timer_state['tiempo_formateado']}")
        else:
            tiempo_restante_placeholder.error(
                f"⏱️ Tiempo agotado: {timer_state['tiempo_formateado']}")

        # Mostrar la tarea
        tarea_simulacro = get_session_var("tarea_simulacro")
        with st.expander("Tarea del simulacro:", expanded=True):
            st.markdown(tarea_simulacro)

        # Área de texto para respuesta
        simulacro_respuesta = st.text_area(
            "Tu respuesta:",
            value=get_session_var("simulacro_respuesta_texto", ""),
            height=300,
            key="simulacro_respuesta_area"
        )

        # Guardar respuesta en tiempo real
        set_session_var("simulacro_respuesta_texto", simulacro_respuesta)

        # Opciones para finalizar o reiniciar
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Finalizar y corregir", key="finalizar_simulacro"):
                if simulacro_respuesta.strip():
                    # Calcular tiempo usado
                    tiempo_final = time.time() - inicio_simulacro
                    minutos_usados = int(tiempo_final // 60)
                    segundos_usados = int(tiempo_final % 60)
                    tiempo_usado = f"{minutos_usados:02d}:{segundos_usados:02d}"

                    # Mostrar spinner durante la corrección
                    with st.spinner("Analizando respuesta..."):
                        # Corrección integrada (sin redirección)
                        resultado = corregir_examen(
                            simulacro_respuesta,
                            tipo_examen,
                            nivel_examen,
                            tiempo_usado
                        )

                    # Limpiar variables de control
                    set_session_var("inicio_simulacro", None)
                    set_session_var("duracion_simulacro", None)
                    set_session_var("tarea_simulacro", None)
                    set_session_var("simulacro_respuesta_texto", "")

                    # Mensaje de éxito
                    st.success(f"Simulacro completado en {tiempo_usado}.")

                    # Mostrar resultados directamente aquí
                    if "error" not in resultado:
                        ui_show_correction_results(resultado)
                    else:
                        st.error(
                            f"Error en la corrección: {resultado['error']}")
                else:
                    st.warning(
                        "Por favor, escribe una respuesta antes de finalizar.")

        with col2:
            if st.button("Reiniciar simulacro", key="reiniciar_simulacro"):
                # Limpiar todas las variables del simulacro
                set_session_var("inicio_simulacro", None)
                set_session_var("duracion_simulacro", None)
                set_session_var("tarea_simulacro", None)
                set_session_var("simulacro_respuesta_texto", "")
                st.rerun()

        # Verificar si se acabó el tiempo
        if timer_state["terminado"]:
            st.error("⏰ ¡Tiempo agotado! Finaliza tu respuesta y envíala.")
            # Guardar automáticamente (opcional)
            st.info(
                "Tu respuesta ha sido guardada automáticamente. Puedes finalizarla ahora.")


def criterios_evaluacion_tab(tipo_examen, nivel_examen):
    """
    Implementación de la pestaña de criterios de evaluación.

    Args:
        tipo_examen: Tipo de examen seleccionado
        nivel_examen: Nivel de examen seleccionado
    """
    st.subheader("Criterios de evaluación")
    st.markdown("""
    Conocer cómo se evalúa tu texto es fundamental para prepararte adecuadamente.
    Aquí encontrarás las rúbricas oficiales y ejemplos de textos evaluados.
    """)

    # Mostrar los criterios específicos según el examen seleccionado
    criterios = obtener_criterios_evaluacion(tipo_examen, nivel_examen)
    st.markdown(criterios)

    # Opción para ver ejemplos evaluados
    if st.button("Ver ejemplos de textos evaluados", key="ver_ejemplos_evaluados"):
        with st.spinner("Generando ejemplos..."):
            ejemplos = generar_ejemplos_evaluados(tipo_examen, nivel_examen)
            st.markdown(ejemplos)
            """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 8: Componentes de UI Avanzados y Pestañas (continuación 5)
==================================================================================

Continuación de los componentes de UI avanzados y pestañas
"""

# --- 5. IMPLEMENTACIÓN DE SUBPESTAÑAS DE HERRAMIENTAS ---


def herramienta_analisis_complejidad():
    """Implementación de la herramienta de análisis de complejidad textual."""
    st.subheader("Análisis de complejidad textual")
    st.markdown("""
    Esta herramienta analiza la complejidad léxica, sintáctica y estructural de tu texto 
    para ayudarte a entender tu nivel actual y cómo mejorar.
    """)

    # Código para el análisis de complejidad
    texto_analisis = st.text_area(
        "Ingresa el texto a analizar:",
        height=200,
        key="texto_analisis"
    )

    if st.button("Analizar complejidad", key="analizar_complejidad") and texto_analisis.strip():
        with st.spinner("Analizando la complejidad de tu texto..."):
            analisis_data = analizar_complejidad_texto(texto_analisis)

            if "error" in analisis_data:
                st.error(f"Error al analizar: {analisis_data['error']}")
                return

            # Mostrar resultados
            st.subheader("Resultados del análisis")

            # Nivel MCER estimado
            nivel_mcer = analisis_data.get("nivel_mcer", {})
            st.info(
                f"📊 **Nivel MCER estimado: {nivel_mcer.get('nivel', 'No disponible')}**")
            st.write(nivel_mcer.get("justificacion", ""))

            # Métricas principales en columnas
            indices = analisis_data.get("indices", {})
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("TTR", f"{indices.get('ttr', 0):.2f}")
                st.caption("Ratio tipo/token - variedad léxica")
            with col2:
                st.metric("Densidad léxica",
                          f"{indices.get('densidad_lexica', 0):.2f}")
                st.caption("Proporción palabras contenido/total")
            with col3:
                st.metric("Índice Szigriszt",
                          f"{indices.get('szigriszt', 0):.1f}")
                st.caption("Legibilidad (70-80: estándar)")

            # Interpretación general
            st.markdown(
                f"**Interpretación general**: {indices.get('interpretacion', '')}")

            # Detalles por áreas
            tabs = st.tabs(["Léxico", "Sintaxis", "Textual"])

            with tabs[0]:
                lex = analisis_data.get("complejidad_lexica", {})
                st.markdown(
                    f"**Nivel de complejidad léxica**: {lex.get('nivel', '')}")
                st.write(lex.get("descripcion", ""))

                palabras = lex.get("palabras_destacadas", [])
                if palabras:
                    st.markdown("**Palabras destacadas:**")
                    st.write(", ".join(palabras))

            with tabs[1]:
                sint = analisis_data.get("complejidad_sintactica", {})
                st.markdown(
                    f"**Nivel de complejidad sintáctica**: {sint.get('nivel', '')}")
                st.write(sint.get("descripcion", ""))

                estructuras = sint.get("estructuras_destacadas", [])
                if estructuras:
                    st.markdown("**Estructuras destacadas:**")
                    for est in estructuras:
                        st.markdown(f"- {est}")

            with tabs[2]:
                text = analisis_data.get("complejidad_textual", {})
                st.markdown(
                    f"**Nivel de complejidad textual**: {text.get('nivel', '')}")
                st.write(text.get("descripcion", ""))

            # Recomendaciones
            recomendaciones = analisis_data.get("recomendaciones", [])
            if recomendaciones:
                with st.expander("Recomendaciones para mejorar", expanded=True):
                    for rec in recomendaciones:
                        st.markdown(f"- {rec}")


def herramienta_biblioteca_recursos():
    """Implementación de la herramienta de biblioteca de recursos."""
    st.subheader("Biblioteca de recursos")
    st.markdown("""
    Accede a recursos didácticos para mejorar tu español, 
    organizados por nivel y categoría gramatical.
    """)

    # Organización de recursos en categorías
    col1, col2 = st.columns(2)

    with col1:
        categoria = st.selectbox(
            "Categoría:",
            [
                "Gramática", "Vocabulario", "Expresiones",
                "Ortografía", "Conectores", "Cultura"
            ],
            key="categoria_recursos"
        )

    with col2:
        nivel_recursos = st.selectbox(
            "Nivel:",
            ["A1", "A2", "B1", "B2", "C1", "C2", "Todos los niveles"],
            key="nivel_recursos"
        )

    # Mapear al nivel en la base de datos
    if nivel_recursos in ["A1", "A2"]:
        nivel_db = "A1-A2"
    elif nivel_recursos in ["B1", "B2"]:
        nivel_db = "B1-B2"
    elif nivel_recursos in ["C1", "C2"]:
        nivel_db = "C1-C2"
    else:
        nivel_db = None  # Todos los niveles

    # Generar recursos basados en la selección
    if st.button("Buscar recursos", key="buscar_recursos"):
        recursos_mostrados = []

        # Buscar en la base de datos estática
        if nivel_db:
            # Filtramos por nivel específico
            nivel_recursos_db = RECURSOS_DB.get(nivel_db, {})
            for cat, recursos in nivel_recursos_db.items():
                if categoria.lower() in cat.lower() or "todos" in categoria.lower():
                    recursos_mostrados.extend(recursos)
        else:
            # Mostrar todos los niveles
            for nivel, categorias in RECURSOS_DB.items():
                for cat, recursos in categorias.items():
                    if categoria.lower() in cat.lower() or "todos" in categoria.lower():
                        recursos_mostrados.extend(recursos)

        # Si no hay recursos, generar con IA
        if not recursos_mostrados:
            with st.spinner("Generando recomendaciones de recursos..."):
                # Generar recursos con OpenAI
                client = get_openai_client()
                if client and circuit_breaker.can_execute("openai"):
                    try:
                        nivel_str = nivel_recursos if nivel_recursos != "Todos los niveles" else "todos los niveles"

                        prompt_recursos = f"""
                        Genera una lista de 5 recursos didácticos reales y relevantes para estudiantes de español 
                        de nivel {nivel_str} enfocados en {categoria}.
                        
                        Cada recurso debe incluir:
                        1. Título descriptivo
                        2. Tipo de recurso (libro, página web, app, podcast, vídeo, etc.)
                        3. URL real (o editorial en caso de libros)
                        4. Breve descripción de su contenido y utilidad
                        5. Nivel específico (si aplica)
                        
                        Devuelve SOLO la información en formato JSON con la estructura:
                        {{
                          "recursos": [
                            {{
                              "titulo": "string",
                              "tipo": "string",
                              "url": "string",
                              "descripcion": "string",
                              "nivel": "string"
                            }}
                          ]
                        }}
                        """

                        response = client.chat.completions.create(
                            model="gpt-4-turbo",
                            temperature=0.5,
                            response_format={"type": "json_object"},
                            messages=[
                                {"role": "system", "content": "Eres un especialista en recursos didácticos para aprendizaje de español como lengua extranjera."},
                                {"role": "user", "content": prompt_recursos}
                            ]
                        )

                        # Extraer JSON
                        content = response.choices[0].message.content
                        recursos_data = extract_json_safely(content)

                        # Verificar si se obtuvo un resultado válido
                        if "error" not in recursos_data:
                            recursos_ia = recursos_data.get("recursos", [])

                            # Convertir al formato de nuestros recursos
                            for recurso in recursos_ia:
                                recursos_mostrados.append({
                                    "título": recurso.get("titulo", ""),
                                    "tipo": recurso.get("tipo", ""),
                                    "url": recurso.get("url", ""),
                                    "nivel": recurso.get("nivel", ""),
                                    "descripcion": recurso.get("descripcion", "")
                                })

                            # Registrar éxito
                            circuit_breaker.record_success("openai")
                        else:
                            st.error(
                                f"Error al generar recursos: {recursos_data['error']}")
                    except Exception as e:
                        st.error(f"Error al generar recursos: {str(e)}")
                        circuit_breaker.record_failure("openai")
                else:
                    st.warning(
                        "Servicio de generación de recursos no disponible en este momento.")

        # Mostrar los recursos
        if recursos_mostrados:
            st.subheader(
                f"Recursos de {categoria} para nivel {nivel_recursos}")

            for i, recurso in enumerate(recursos_mostrados):
                with st.expander(f"{i+1}. {recurso.get('título', '')} ({recurso.get('nivel', '')})", expanded=i == 0):
                    st.markdown(f"**Tipo:** {recurso.get('tipo', '')}")
                    st.markdown(
                        f"**URL:** [{recurso.get('url', '').split('/')[-1]}]({recurso.get('url', '')})")
                    if "descripcion" in recurso:
                        st.markdown(
                            f"**Descripción:** {recurso.get('descripcion', '')}")
        else:
            st.info(
                f"No se encontraron recursos para {categoria} de nivel {nivel_recursos}. Intenta con otra combinación.")
            """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 8: Componentes de UI Avanzados y Pestañas (continuación 6)
==================================================================================

Continuación de los componentes de UI avanzados y pestañas
"""


def herramienta_descripcion_imagenes():
    """Implementación de la herramienta de descripción de imágenes con DALL-E."""
    st.subheader("🖼️ Descripción de imágenes generadas por IA")
    st.markdown("""
    Esta herramienta genera imágenes adaptadas a tu nivel de español y proporciona actividades
    de descripción para practicar vocabulario y estructuras descriptivas.
    """)

    # Usamos un enfoque diferente para mantener el estado, sin modificar directamente
    # las claves de session_state vinculadas a widgets
    if "tab_navigate_to" not in st.session_state:
        st.session_state.tab_navigate_to = None

    # Establecer las pestañas correctas para futuras navegaciones
    # Esto sólo afecta a futuras recargas de la página, no a la actual
    st.session_state.tab_navigate_to = {
        "main_tab": 4,  # Índice de "Herramientas complementarias"
        "tools_tab": 2  # Índice de "Descripción de imágenes"
    }

    # Selección de nivel
    nivel_imagen = st.selectbox(
        "Nivel de español:",
        [
            "Nivel principiante (A1-A2)",
            "Nivel intermedio (B1-B2)",
            "Nivel avanzado (C1-C2)"
        ],
        index=["principiante", "intermedio", "avanzado"].index(
            get_session_var("nivel_estudiante", "intermedio")
        ),
        key="nivel_imagen_dalle"
    )

    # Tema para la imagen
    tema_imagen = st.text_input(
        "Tema o escena para la imagen (por ejemplo: 'un parque en primavera', 'una oficina moderna'):",
        key="tema_imagen_dalle",
        value=st.session_state.tema_imagen_state if st.session_state.tema_imagen_state else ""
    )

    # Función para generar la imagen
    def generar_imagen_callback():
        if not tema_imagen:
            st.warning("Por favor, introduce un tema para la imagen.")
            return

        with st.spinner("Generando imagen con DALL-E..."):
            # Obtener nivel en formato simplificado
            nivel_map = {
                "Nivel principiante (A1-A2)": "principiante",
                "Nivel intermedio (B1-B2)": "intermedio",
                "Nivel avanzado (C1-C2)": "avanzado"
            }
            nivel_dalle = nivel_map.get(nivel_imagen, "intermedio")

            # Generar imagen y descripción
            imagen_url, descripcion = generar_imagen_dalle(
                tema_imagen, nivel_dalle)

            if imagen_url:
                # Guardar en session_state
                st.session_state.imagen_generada_state = True
                st.session_state.imagen_url_state = imagen_url
                st.session_state.descripcion_state = descripcion
                st.session_state.tema_imagen_state = tema_imagen

                # También guardar en las variables originales
                set_session_var("ultima_imagen_url", imagen_url)
                set_session_var("ultima_descripcion", descripcion)

                # Asegurar que estamos en las pestañas correctas para la próxima recarga
                if "tab_navigate_to" not in st.session_state:
                    st.session_state.tab_navigate_to = {}
                st.session_state.tab_navigate_to = {
                    "main_tab": 4,  # Herramientas complementarias
                    "tools_tab": 2   # Descripción de imágenes
                }

                # Forzar navegación a las pestañas correctas en el inicio de la app
                st.session_state.active_tab_index = 4
                st.session_state.active_tools_tab_index = 2

                st.success("Imagen generada con éxito")
                # No usamos st.rerun() aquí para evitar problemas con session_state
            else:
                st.error(
                    "No se pudo generar la imagen. Por favor, inténtalo de nuevo.")

    # Botón para generar imagen
    if st.button("Generar imagen y actividad", key="generar_imagen_dalle", on_click=generar_imagen_callback):
        pass  # La lógica se maneja en el callback

    # Mostrar la imagen si existe
    if st.session_state.imagen_generada_state and st.session_state.imagen_url_state:
        # Mostrar la imagen
        st.image(
            st.session_state.imagen_url_state,
            caption=f"Imagen generada sobre: {st.session_state.tema_imagen_state}",
            use_container_width=True
        )

        # Mostrar la descripción y actividades
        with st.expander("Descripción y actividades de práctica", expanded=True):
            st.markdown(st.session_state.descripcion_state)

        # Área para que el estudiante escriba su descripción
        descripcion_estudiante = st.text_area(
            "Describe la imagen con tus propias palabras:",
            height=200,
            key="descripcion_imagen_estudiante",
            value=st.session_state.descripcion_estudiante_state
        )

        # Actualizar el valor en la sesión cuando cambie
        st.session_state.descripcion_estudiante_state = descripcion_estudiante

        # Función para corregir la descripción
        def corregir_descripcion_callback():
            if not descripcion_estudiante.strip():
                st.warning(
                    "Por favor, escribe una descripción antes de enviar a corrección.")
                return

            # Asegurar que estamos en las pestañas correctas para la próxima recarga
            st.session_state.active_tab_index = 4
            st.session_state.active_tools_tab_index = 2

            # Indicar que queremos mostrar la corrección
            st.session_state.mostrar_correccion_imagen = True

        # Botón para enviar a corrección
        if st.button("Corregir descripción", key="corregir_descripcion_imagen", on_click=corregir_descripcion_callback):
            pass  # La lógica se maneja en el callback

        # Mostrar resultados de corrección si es necesario
        if st.session_state.get("mostrar_correccion_imagen", False):
            with st.spinner("Analizando descripción..."):
                # Corrección integrada
                resultado = corregir_descripcion_imagen(
                    descripcion_estudiante,
                    st.session_state.tema_imagen_state,
                    nivel_imagen
                )

                # Mostrar resultados
                if "error" not in resultado:
                    # Mostrar la corrección
                    ui_show_correction_results(resultado)
                else:
                    st.error(f"Error en la corrección: {resultado['error']}")

                # Opción para volver a la imagen
                if st.button("Volver a la descripción de imagen", key="volver_a_imagen"):
                    st.session_state.mostrar_correccion_imagen = False
                    st.experimental_rerun()

    # Botón para reiniciar
    if st.session_state.imagen_generada_state:
        if st.button("Generar nueva imagen", key="nueva_imagen_dalle"):
            # Reiniciar estados
            st.session_state.imagen_generada_state = False
            st.session_state.imagen_url_state = None
            st.session_state.descripcion_state = None
            st.session_state.tema_imagen_state = None
            st.session_state.descripcion_estudiante_state = ""
            st.session_state.mostrar_correccion_imagen = False
            st.experimental_rerun()


def herramienta_texto_manuscrito():
    """
    Implementación corregida de la herramienta de transcripción de textos manuscritos.
    Soluciona problemas de flujo entre transcripción y corrección.
    """
    st.subheader("✍️ Transcripción de textos manuscritos")
    st.markdown("""
    Esta herramienta te permite subir imágenes de textos manuscritos para transcribirlos
    automáticamente y luego enviarlos a corrección.
    """)

    # IMPORTANTE: Verificar si estamos en modo de corrección de transcripción
    if get_session_var("mostrar_correccion_transcripcion", False):
        visualizar_texto_manuscrito()
        return

    # Selección de idioma para la transcripción
    idioma_manuscrito = st.selectbox(
        "Idioma del texto manuscrito:",
        ["Español", "Francés", "Inglés"],
        key="idioma_manuscrito"
    )

    # Mapeo de idiomas para la API
    idioma_map = {
        "Español": "es",
        "Francés": "fr",
        "Inglés": "en"
    }

    # Subida de imagen
    imagen_manuscrito = st.file_uploader(
        "Sube una imagen de tu texto manuscrito (JPG, PNG):",
        type=["jpg", "jpeg", "png"],
        key="imagen_manuscrito"
    )

    if imagen_manuscrito is not None:
        # Mostrar la imagen subida
        try:
            imagen = Image.open(imagen_manuscrito)
            st.image(imagen, caption="Imagen subida", use_column_width=True)
        except Exception as e:
            st.error(f"Error al procesar la imagen: {str(e)}")
            return

        # Botón para transcribir
        if st.button("Transcribir texto", key="transcribir_manuscrito"):
            with st.spinner("Transcribiendo texto manuscrito..."):
                try:
                    # Leer bytes de la imagen
                    imagen_bytes = imagen_manuscrito.getvalue()

                    # Obtener código de idioma
                    codigo_idioma = idioma_map.get(idioma_manuscrito, "es")

                    # Transcribir la imagen
                    texto_transcrito = transcribir_imagen_texto(
                        imagen_bytes, codigo_idioma
                    )

                    if texto_transcrito and not texto_transcrito.startswith("Error"):
                        # Mostrar el texto transcrito
                        st.success("✅ Texto transcrito correctamente")

                        with st.expander("Texto transcrito", expanded=True):
                            st.write(texto_transcrito)

                            # Guardar en session_state de forma segura
                            set_session_var(
                                "ultimo_texto_transcrito", texto_transcrito)

                        # CORRECCIÓN CLAVE: En lugar de redireccionar, usar una bandera
                        if st.button("Corregir texto transcrito", key="corregir_texto_transcrito"):
                            set_session_var(
                                "mostrar_correccion_transcripcion", True)
                            st.rerun()  # Forzar actualización de la interfaz
                    else:
                        st.error(
                            texto_transcrito or "No se pudo transcribir el texto. Por favor, verifica que la imagen sea clara y contiene texto manuscrito legible.")
                except Exception as e:
                    st.error(f"Error durante la transcripción: {str(e)}")
                    logger.error(
                        f"Error en herramienta_texto_manuscrito: {str(e)}")
                    logger.error(traceback.format_exc())
                    """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 8: Componentes de UI Avanzados y Pestañas (continuación 7)
==================================================================================

Continuación de los componentes de UI avanzados y pestañas
"""

# --- 6. IMPLEMENTACIÓN DE PESTAÑAS DE PROGRESO Y ESTADÍSTICAS ---


def estadisticas_progreso_tab():
    """Implementación de la pestaña de estadísticas de progreso."""
    nombre_estudiante = st.text_input(
        "Nombre y apellido del estudiante para ver progreso:",
        value=get_session_var("usuario_actual", ""),
        key="nombre_progreso"
    )

    if nombre_estudiante and " " not in nombre_estudiante:
        st.warning(
            "Por favor, introduce tanto el nombre como el apellido separados por un espacio.")

    if nombre_estudiante and " " in nombre_estudiante:
        with st.spinner("Cargando datos de progreso..."):
            try:
                df = obtener_historial_estudiante(nombre_estudiante)
                if df is not None and not df.empty:
                    # Generar gráficos
                    graficos = mostrar_progreso(df)

                    # Mostrar gráficos si existen
                    if graficos["errores_totales"] is not None:
                        st.subheader("Progreso en la reducción de errores")
                        st.altair_chart(
                            graficos["errores_totales"], use_container_width=True)

                    if graficos["tipos_error"] is not None:
                        st.altair_chart(
                            graficos["tipos_error"], use_container_width=True)

                    if graficos["radar"] is not None:
                        st.pyplot(graficos["radar"])

                    # Mostrar tabla con historial completo
                    with st.expander("Ver datos completos"):
                        st.dataframe(df)

                    # Verificar si hay suficientes datos para análisis de tendencias
                    if len(df) >= 2 and graficos["fecha_col"] is not None:
                        fecha_col = graficos["fecha_col"]

                        # Consejo basado en tendencias
                        st.subheader("Consejo basado en tendencias")

                        # Calcular tendencias simples
                        df[fecha_col] = pd.to_datetime(df[fecha_col])
                        df = df.sort_values(fecha_col)

                        # Extraer primera y última entrada para comparar
                        primera = df.iloc[0]
                        ultima = df.iloc[-1]

                        # Comparar total de errores - asegurar que son numéricos
                        primera_errores = float(
                            primera.get('Total Errores', 0))
                        ultima_errores = float(ultima.get('Total Errores', 0))
                        dif_errores = ultima_errores - primera_errores

                        if dif_errores < 0:
                            st.success(
                                f"¡Felicidades! Has reducido tus errores en {abs(dif_errores)} desde tu primera entrega.")
                        elif dif_errores > 0:
                            st.warning(
                                f"Has aumentado tus errores en {dif_errores} desde tu primera entrega. Revisa las recomendaciones.")
                        else:
                            st.info(
                                "El número total de errores se mantiene igual. Sigamos trabajando en las áreas de mejora.")

                        # Identificar área con mayor progreso y área que necesita más trabajo
                        categorias = [
                            'Errores Gramática', 'Errores Léxico', 'Errores Puntuación', 'Errores Estructura']
                        categorias_existentes = [
                            cat for cat in categorias if cat in df.columns]

                        if categorias_existentes:
                            difs = {}
                            for cat in categorias_existentes:
                                if cat in primera and cat in ultima:
                                    # Asegurar conversión a numérico
                                    difs[cat] = float(ultima.get(
                                        cat, 0)) - float(primera.get(cat, 0))

                            if difs:
                                mejor_area = min(
                                    difs.items(), key=lambda x: x[1])[0]
                                peor_area = max(
                                    difs.items(), key=lambda x: x[1])[0]

                                if difs[mejor_area] < 0:
                                    st.success(
                                        f"Mayor progreso en: {mejor_area.replace('Errores ', '')}")

                                if difs[peor_area] > 0:
                                    st.warning(
                                        f"Área que necesita más trabajo: {peor_area.replace('Errores ', '')}")
                else:
                    st.info(
                        f"No se encontraron datos para '{nombre_estudiante}' en el historial.")

                    # Mostrar nombres disponibles
                    if sheets_connection is not None and sheets_connection["tracking"] is not None:
                        try:
                            todos_datos = sheets_connection["tracking"].get_all_records(
                            )
                            if todos_datos:
                                columnas = list(todos_datos[0].keys())
                                nombre_col = next(
                                    (col for col in columnas if col.lower() == 'nombre'), None)

                                if nombre_col:
                                    nombres_disponibles = sorted(set(str(row.get(nombre_col, '')).strip()
                                                                     for row in todos_datos if row.get(nombre_col)))

                                    if nombres_disponibles:
                                        st.write(
                                            "Nombres disponibles en el historial:")

                                        # Dividir en filas de 3 botones
                                        for i in range(0, len(nombres_disponibles), 3):
                                            fila = nombres_disponibles[i:i+3]
                                            cols = st.columns(3)
                                            for j, nombre in enumerate(fila):
                                                if j < len(fila) and cols[j].button(nombre, key=f"btn_progreso_{nombre}_{i+j}"):
                                                    # Establecer el nombre seleccionado
                                                    set_session_var(
                                                        "nombre_seleccionado", nombre)
                                                    st.rerun()
                        except Exception as e:
                            st.error(
                                f"Error al listar nombres disponibles: {str(e)}")
            except Exception as e:
                st.error(f"Error al obtener historial: {str(e)}")
                with st.expander("Detalles del error"):
                    st.code(traceback.format_exc())


def plan_estudio_tab():
    """Implementación de la pestaña de plan de estudio personalizado."""
    st.header("📚 Plan de estudio personalizado")

    nombre_estudiante_plan = st.text_input(
        "Nombre y apellido:",
        value=get_session_var("usuario_actual", ""),
        key="nombre_plan_estudio"
    )

    if nombre_estudiante_plan and " " not in nombre_estudiante_plan:
        st.warning(
            "Por favor, introduce tanto el nombre como el apellido separados por un espacio.")

    if nombre_estudiante_plan and " " in nombre_estudiante_plan:
        with st.spinner("Analizando tu historial de errores y generando plan personalizado..."):
            # Obtener historial del estudiante
            df = obtener_historial_estudiante(nombre_estudiante_plan)

            if df is not None and not df.empty:
                # Generar plan de estudio personalizado
                resultado = generar_plan_estudio_personalizado(
                    nombre_estudiante_plan,
                    get_session_var("nivel_estudiante", "intermedio"),
                    df
                )

                if resultado["error"] is None and resultado["plan"] is not None:
                    plan = resultado["plan"]

                    st.markdown("### Tu plan de estudio personalizado")
                    st.markdown(
                        "Basado en tu historial de errores, hemos creado este plan de estudio de 4 semanas para ayudarte a mejorar tus habilidades:")

                    # Mostrar cada semana en un expander
                    for semana in plan["semanas"]:
                        with st.expander(f"Semana {semana['numero']}: {semana['titulo']}", expanded=semana['numero'] == 1):
                            st.markdown(semana['contenido'])

                            # Generar ejercicios específicos para esta parte
                            if st.button(f"Generar ejercicios para Semana {semana['numero']}", key=f"ejercicios_semana_{semana['numero']}"):
                                with st.spinner("Creando ejercicios personalizados..."):
                                    client = get_openai_client()
                                    if client and circuit_breaker.can_execute("openai"):
                                        try:
                                            prompt_ejercicios = f"Crea 2 ejercicios breves para practicar los temas de la semana {semana['numero']} del plan: {semana['contenido'][:300]}... Los ejercicios deben ser específicos para un estudiante de nivel {get_session_var('nivel_estudiante', 'intermedio')}."

                                            response = client.chat.completions.create(
                                                model="gpt-4-turbo",
                                                temperature=0.7,
                                                messages=[
                                                    {"role": "system", "content": "Eres un profesor de español especializado en crear actividades didácticas."},
                                                    {"role": "user",
                                                        "content": prompt_ejercicios}
                                                ]
                                            )

                                            ejercicios = response.choices[0].message.content
                                            st.markdown(
                                                "#### Ejercicios recomendados")
                                            st.markdown(ejercicios)

                                            # Registrar éxito
                                            circuit_breaker.record_success(
                                                "openai")
                                        except Exception as e:
                                            st.error(
                                                f"Error al generar ejercicios: {str(e)}")
                                            circuit_breaker.record_failure(
                                                "openai")
                                    else:
                                        st.warning(
                                            "Servicio de generación de ejercicios no disponible en este momento.")
                else:
                    st.error(resultado["error"])
            else:
                st.info("No tenemos suficientes datos para generar un plan personalizado. Realiza al menos 3 correcciones de texto para activar esta función.")
                """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 8: Componentes de UI Avanzados y Pestañas (continuación 8)
==================================================================================

Continuación de los componentes de UI avanzados y pestañas
"""

# --- 7. IMPLEMENTACIÓN DE PESTAÑAS PRINCIPALES ---


def tab_progreso():
    """Implementación de la pestaña de progreso."""
    # Limpiar variables de simulacro si venimos de otra pestaña
    if "inicio_simulacro" in st.session_state:
        set_session_var("inicio_simulacro", None)
    if "duracion_simulacro" in st.session_state:
        set_session_var("duracion_simulacro", None)
    if "tarea_simulacro" in st.session_state:
        set_session_var("tarea_simulacro", None)

    st.header("📊 Seguimiento del progreso")

    # Verificar si hay usuario
    if "usuario_actual" not in st.session_state or not st.session_state.usuario_actual:
        st.info(
            "👆 Por favor, introduce tu nombre y nivel en la pestaña 'Corrección de texto' para comenzar.")

        # Mostrar formulario básico de usuario con key única para este contexto
        user_data = ui_user_info_form(form_key="form_user_info_progreso")
        if not user_data:
            return

    # Subtabs para diferentes vistas de progreso
    subtab_estadisticas, subtab_plan_estudio = st.tabs([
        "Estadísticas", "Plan de estudio personalizado"
    ])

    with subtab_estadisticas:
        estadisticas_progreso_tab()

    with subtab_plan_estudio:
        plan_estudio_tab()


def tab_historial():
    """Implementación de la pestaña de historial."""
    # Limpiar variables de simulacro si venimos de otra pestaña
    if "inicio_simulacro" in st.session_state:
        set_session_var("inicio_simulacro", None)
    if "duracion_simulacro" in st.session_state:
        set_session_var("duracion_simulacro", None)
    if "tarea_simulacro" in st.session_state:
        set_session_var("tarea_simulacro", None)

    st.header("📚 Historial de correcciones")

    # Verificar si hay usuario
    if "usuario_actual" not in st.session_state or not st.session_state.usuario_actual:
        st.info(
            "👆 Por favor, introduce tu nombre y nivel en la pestaña 'Corrección de texto' para comenzar.")

        # Mostrar formulario básico de usuario con key única para este contexto
        user_data = ui_user_info_form(form_key="form_user_info_historial")
        if not user_data:
            return

    if sheets_connection is None or sheets_connection["corrections"] is None:
        st.warning("⚠️ No hay conexión con la base de datos de correcciones.")
        return

    try:
        # Obtener todas las correcciones
        correcciones = sheets_connection["corrections"].get_all_records()

        if correcciones:
            # Convertir a dataframe
            df_correcciones = pd.DataFrame(correcciones)

            # Normalizar nombres de columnas para la verificación (convertir a minúsculas)
            df_columns_lower = [col.lower() for col in df_correcciones.columns]

            # Filtrar columnas relevantes (verificando de forma más flexible)
            if 'nombre' in df_columns_lower or 'Nombre' in df_correcciones.columns:
                # Determinar los nombres reales de las columnas
                nombre_col = 'Nombre' if 'Nombre' in df_correcciones.columns else 'nombre'
                nivel_col = 'Nivel' if 'Nivel' in df_correcciones.columns else 'nivel'
                fecha_col = 'Fecha' if 'Fecha' in df_correcciones.columns else 'fecha'

                # Verificar que todas las columnas existan
                if nombre_col in df_correcciones.columns and nivel_col in df_correcciones.columns and fecha_col in df_correcciones.columns:
                    df_display = df_correcciones[[
                        nombre_col, nivel_col, fecha_col]]

                    # Mostrar tabla de historial con filtro
                    st.subheader("Correcciones guardadas")

                    # Filtro por nombre
                    nombres_unicos = sorted(
                        df_correcciones[nombre_col].unique().tolist())
                    nombre_filtro = st.selectbox(
                        "Filtrar por nombre:",
                        ["Todos"] + nombres_unicos,
                        index=0,
                        key="nombre_filtro_historial"
                    )

                    # Aplicar filtro
                    if nombre_filtro != "Todos":
                        df_filtrado = df_display[df_display[nombre_col]
                                                 == nombre_filtro]
                    else:
                        df_filtrado = df_display

                    # Mostrar datos filtrados
                    st.dataframe(df_filtrado)

                    # Opciones para ver detalles
                    if st.checkbox("Ver detalles de una corrección", key="checkbox_historial"):
                        # Extraer nombres únicos
                        nombres = sorted(
                            df_correcciones[nombre_col].unique().tolist())

                        # Selector de nombre
                        nombre_select = st.selectbox(
                            "Selecciona un nombre:",
                            nombres,
                            key="nombre_select_historial"
                        )

                        # Filtrar por nombre
                        correcciones_filtradas = df_correcciones[df_correcciones[nombre_col]
                                                                 == nombre_select]

                        # Extraer fechas para este nombre
                        fechas = correcciones_filtradas[fecha_col].tolist()

                        # Selector de fecha
                        fecha_select = st.selectbox(
                            "Selecciona una fecha:",
                            fechas,
                            key="fecha_select_historial"
                        )

                        # Mostrar corrección seleccionada
                        correccion = correcciones_filtradas[correcciones_filtradas[fecha_col]
                                                            == fecha_select].iloc[0]

                        # Mostrar detalles
                        st.subheader(
                            f"Corrección para {nombre_select} ({fecha_select})")

                        # Pestañas para texto original y datos
                        tab_original, tab_datos = st.tabs(
                            ["Texto original", "Datos de corrección"]
                        )

                        with tab_original:
                            texto_col = 'texto' if 'texto' in df_correcciones.columns else 'Texto'
                            if texto_col in correccion:
                                st.write(correccion.get(
                                    texto_col, 'No disponible'))
                            else:
                                st.warning(
                                    "No se pudo encontrar el texto original.")

                        with tab_datos:
                            try:
                                # Intentar parsear el JSON de la respuesta
                                raw_output_col = 'raw_output' if 'raw_output' in df_correcciones.columns else 'Raw_output'
                                if raw_output_col in correccion:
                                    raw_output = correccion.get(
                                        raw_output_col, '{}')
                                    data_json = extract_json_safely(raw_output)

                                    # Verificar si se obtuvo un resultado válido
                                    if "error" not in data_json:
                                        # Mostrar campos específicos
                                        if 'texto_corregido' in data_json:
                                            st.subheader("Texto corregido")
                                            st.write(
                                                data_json['texto_corregido'])

                                        if 'consejo_final' in data_json:
                                            st.subheader("Consejo final")
                                            st.info(data_json['consejo_final'])
                                    else:
                                        st.warning(
                                            "No se pudieron cargar los datos de corrección en formato estructurado.")
                                        # Mostrar parte del texto crudo
                                        st.code(raw_output[:500] + "...")
                                else:
                                    st.warning(
                                        "No se encontraron datos de corrección.")
                            except Exception as e:
                                st.warning(
                                    "No se pudieron cargar los datos de corrección en formato estructurado.")
                                # Mostrar parte del texto crudo como fallback
                                st.code(
                                    raw_output_col[:500] + "..." if raw_output_col else "Datos no disponibles")
                else:
                    st.warning("Algunas columnas necesarias no se encuentran en los datos. Columnas disponibles: " +
                               ", ".join(df_correcciones.columns))
            else:
                st.warning("El formato de los datos no coincide con lo esperado. Columnas disponibles: " +
                           ", ".join(df_correcciones.columns))
        else:
            st.info("No hay correcciones guardadas en el historial.")
    except Exception as e:
        st.error(f"Error al cargar el historial: {str(e)}")
        with st.expander("Detalles del error"):
            st.code(traceback.format_exc())


def tab_examenes():
    """Implementación de la pestaña de exámenes."""
    st.header("🎓 Preparación para exámenes oficiales")

    # Verificar si hay usuario
    if "usuario_actual" not in st.session_state or not st.session_state.usuario_actual:
        st.info(
            "👆 Por favor, introduce tu nombre y nivel en la pestaña 'Corrección de texto' para comenzar.")

        # Mostrar formulario básico de usuario con key única para este contexto
        user_data = ui_user_info_form(form_key="form_user_info_examenes")
        if not user_data:
            return

    # Selector de examen y nivel
    exam_options = ui_examen_options()
    tipo_examen = exam_options["tipo_examen"]
    nivel_examen = exam_options["nivel_examen"]

    # Pestañas para las diferentes funcionalidades
    tabs_examen = st.tabs([
        "Modelo de examen",
        "Simulacro cronometrado",
        "Criterios de evaluación"
    ])

    # --- Pestaña 1: Modelo de examen ---
    with tabs_examen[0]:
        modelo_examen_tab(tipo_examen, nivel_examen)

    # --- Pestaña 2: Simulacro cronometrado ---
    with tabs_examen[1]:
        simulacro_cronometrado_tab(tipo_examen, nivel_examen)

    # --- Pestaña 3: Criterios de evaluación ---
    with tabs_examen[2]:
        criterios_evaluacion_tab(tipo_examen, nivel_examen)


def tab_herramientas():
    """Implementación de la pestaña de herramientas complementarias."""
    # Limpiar variables de simulacro si venimos de otra pestaña
    if "inicio_simulacro" in st.session_state:
        set_session_var("inicio_simulacro", None)
    if "duracion_simulacro" in st.session_state:
        set_session_var("duracion_simulacro", None)
    if "tarea_simulacro" in st.session_state:
        set_session_var("tarea_simulacro", None)

    st.header("🔧 Herramientas complementarias")

    # Verificar si hay usuario
    if "usuario_actual" not in st.session_state or not st.session_state.usuario_actual:
        st.info(
            "👆 Por favor, introduce tu nombre y nivel en la pestaña 'Corrección de texto' para comenzar.")

        # Mostrar formulario básico de usuario con key única para este contexto
        user_data = ui_user_info_form(form_key="form_user_info_herramientas")
        if not user_data:
            return

    # Función para cambiar la subpestaña activa
    def on_tools_tab_change():
        st.session_state.active_tools_tab_index = st.session_state.tools_tab_selector

    # Pestañas para diferentes herramientas
    tools_tab_names = [
        "Análisis de complejidad",
        "Biblioteca de recursos",
        "Descripción de imágenes",
        "Texto manuscrito"
    ]

    # Selector de herramientas
    selected_tools_tab = st.radio(
        "Herramientas",
        options=range(len(tools_tab_names)),
        format_func=lambda x: tools_tab_names[x],
        key="tools_tab_selector",
        horizontal=True,
        label_visibility="collapsed",
        index=st.session_state.active_tools_tab_index,
        on_change=on_tools_tab_change
    )

    # Actualizar índice de subpestaña activa
    st.session_state.active_tools_tab_index = selected_tools_tab

    # --- Subpestaña 1: Análisis de complejidad ---
    if selected_tools_tab == 0:
        herramienta_analisis_complejidad()

    # --- Subpestaña 2: Biblioteca de recursos ---
    elif selected_tools_tab == 1:
        herramienta_biblioteca_recursos()

    # --- Subpestaña 3: Descripción de imágenes ---
    elif selected_tools_tab == 2:
        herramienta_descripcion_imagenes()

    # --- Subpestaña 4: Texto manuscrito ---
    elif selected_tools_tab == 3:
        herramienta_texto_manuscrito()
        """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 9: Lógica de Navegación y Manejo de Estados
==================================================================================

Este artefacto contiene:
1. Manejo de parámetros de URL y comandos de navegación
2. Inicialización de la aplicación
3. Funciones para manejar mantenimiento y estado global
"""

# --- 1. MANEJO DE COMANDOS DE URL Y PARÁMETROS ---


def handle_url_params_fix():
    """Maneja los parámetros de URL para navegación entre páginas."""
    # Al usar st.query_params, se accede como un dict pero los valores no vienen en listas
    # Actualización para mantener comportamiento compatible con el código existente

    # Verificar si hay parámetros que necesitamos procesar
    if "nombre_seleccionado" in st.query_params:
        nombre = st.query_params["nombre_seleccionado"]
        if nombre:
            set_session_var("usuario_actual", nombre)
        # Eliminar el parámetro después de usarlo
        del st.query_params["nombre_seleccionado"]

    if "tab" in st.query_params:
        tab = st.query_params["tab"]
        st.session_state.active_tab = tab
        # Eliminar el parámetro después de usarlo
        del st.query_params["tab"]

# --- 2. INICIALIZACIÓN DE LA APLICACIÓN ---


def init_app():
    """Inicializa la aplicación con comprobaciones y configuraciones necesarias."""
    # Verificar dependencias externas
    dependencies_ok = True

    # Verificar API keys
    if api_keys["openai"] is None:
        dependencies_ok = False
        logger.warning("OpenAI API key no configurada")

    if api_keys["elevenlabs"]["api_key"] is None or api_keys["elevenlabs"]["voice_id"] is None:
        logger.warning("ElevenLabs no configurado completamente")

    if sheets_connection is None:
        logger.warning("Conexión a Google Sheets no disponible")

    # Configurar el estado de la página
    if not dependencies_ok:
        st.sidebar.warning(
            "⚠️ Algunas funcionalidades están limitadas debido a configuraciones incompletas.")

    # Inicializar índices de pestañas si no existen
    if "active_tab_index" not in st.session_state:
        st.session_state.active_tab_index = 0
    if "active_tools_tab_index" not in st.session_state:
        st.session_state.active_tools_tab_index = 0

    # Comprobar si hay un objetivo de navegación específico
    if "tab_navigate_to" in st.session_state and st.session_state.tab_navigate_to:
        try:
            if "main_tab" in st.session_state.tab_navigate_to:
                st.session_state.active_tab_index = st.session_state.tab_navigate_to["main_tab"]
            if "tools_tab" in st.session_state.tab_navigate_to:
                st.session_state.active_tools_tab_index = st.session_state.tab_navigate_to[
                    "tools_tab"]
            # Limpiar después de usar
            st.session_state.tab_navigate_to = None
        except Exception as e:
            logger.error(f"Error al establecer pestañas activas: {str(e)}")

    # Manejar parámetros de URL
    handle_url_params_fix()

    # Comprobar si es la primera ejecución
    if "app_initialized" not in st.session_state:
        set_session_var("app_initialized", True)
        logger.info(
            f"Aplicación inicializada con ID de sesión: {st.session_state.session_id}")

    # Mostrar información de versión y estado en el pie de página
    with st.sidebar.expander("ℹ️ Acerca de la aplicación"):
        st.write(f"Versión: {APP_VERSION}")
        st.write(f"ID de sesión: {st.session_state.session_id[:8]}")

        # Mostrar estado de conexiones
        st.subheader("Estado de servicios")
        circuit_status = circuit_breaker.get_status()

        for service, status in circuit_status.items():
            if status["open"]:
                st.error(f"❌ {service.title()}: No disponible")
            else:
                st.success(f"✅ {service.title()}: Disponible")
                """
TEXTOCORRECTOR ELE - APLICACIÓN DE CORRECCIÓN DE TEXTOS EN ESPAÑOL CON ANÁLISIS CONTEXTUAL
==================================================================================
Artefacto 10: Aplicación Principal (main)
==================================================================================

Este artefacto contiene:
1. Función principal de la aplicación (main)
2. Punto de entrada de la aplicación
3. Manejo de errores de nivel superior
"""

# --- 1. CONFIGURACIÓN DE LA APLICACIÓN PRINCIPAL ---


def main():
    """Función principal de la aplicación Textocorrector ELE."""
    # Título y descripción
    ui_header()

    # Inicializar el seguimiento de pestañas activas si no existe
    if "active_tab_index" not in st.session_state:
        st.session_state.active_tab_index = 0
    if "active_tools_tab_index" not in st.session_state:
        st.session_state.active_tools_tab_index = 0

    # Pestañas principales
    tab_names = [
        "📝 Corrección de texto",
        "📊 Ver progreso",
        "📚 Historial",
        "🎓 Preparación para exámenes",
        "🔧 Herramientas complementarias"
    ]

    # Función para cambiar la pestaña activa
    def on_tab_change():
        st.session_state.active_tab_index = st.session_state.tab_selector

    # Selector de pestañas como radio buttons
    selected_tab = st.radio(
        "Navegación principal",
        options=range(len(tab_names)),
        format_func=lambda x: tab_names[x],
        key="tab_selector",
        horizontal=True,
        label_visibility="collapsed",
        index=st.session_state.active_tab_index,
        on_change=on_tab_change
    )

    # Actualizar índice de pestaña activa
    st.session_state.active_tab_index = selected_tab

    # --- Pestaña 1: Corrección de Texto ---
    if selected_tab == 0:
        tab_corregir()

    # --- Pestaña 2: Ver Progreso ---
    elif selected_tab == 1:
        tab_progreso()

    # --- Pestaña 3: Historial ---
    elif selected_tab == 2:
        tab_historial()

    # --- Pestaña 4: Preparación para Exámenes ---
    elif selected_tab == 3:
        tab_examenes()

    # --- Pestaña 5: Herramientas Complementarias ---
    elif selected_tab == 4:
        # Implementación de la pestaña de herramientas complementarias
        st.header("🔧 Herramientas complementarias")

        # Verificar si hay usuario
        if "usuario_actual" not in st.session_state or not st.session_state.usuario_actual:
            st.info(
                "👆 Por favor, introduce tu nombre y nivel en la pestaña 'Corrección de texto' para comenzar.")

            # Mostrar formulario básico de usuario con key única para este contexto
            user_data = ui_user_info_form(
                form_key="form_user_info_herramientas")
            if not user_data:
                return

        # Pestañas para diferentes herramientas
        tools_tab_names = [
            "Análisis de complejidad",
            "Biblioteca de recursos",
            "Descripción de imágenes",
            "Texto manuscrito"
        ]

        # Función para cambiar la subpestaña activa
        def on_tools_tab_change():
            st.session_state.active_tools_tab_index = st.session_state.tools_tab_selector

        # Selector de herramientas
        selected_tools_tab = st.radio(
            "Herramientas",
            options=range(len(tools_tab_names)),
            format_func=lambda x: tools_tab_names[x],
            key="tools_tab_selector",
            horizontal=True,
            label_visibility="collapsed",
            index=st.session_state.active_tools_tab_index,
            on_change=on_tools_tab_change
        )

        # Actualizar índice de subpestaña activa
        st.session_state.active_tools_tab_index = selected_tools_tab

        # --- Subpestaña 1: Análisis de complejidad ---
        if selected_tools_tab == 0:
            herramienta_analisis_complejidad()

        # --- Subpestaña 2: Biblioteca de recursos ---
        elif selected_tools_tab == 1:
            herramienta_biblioteca_recursos()

        # --- Subpestaña 3: Descripción de imágenes ---
        elif selected_tools_tab == 2:
            herramienta_descripcion_imagenes()

        # --- Subpestaña 4: Texto manuscrito ---
        elif selected_tools_tab == 3:
            herramienta_texto_manuscrito()

    # Formulario de feedback al final
    ui_feedback_form()


# --- 4. PUNTO DE ENTRADA PRINCIPAL ---
# Llamar a la inicialización
init_app()

# Ejecutar la aplicación principal
if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error("❌ Se ha producido un error en la aplicación.")
        with st.expander("Detalles del error", expanded=False):
            st.code(traceback.format_exc())

        logger.error(f"Error en la aplicación: {str(e)}")
        logger.error(traceback.format_exc())
